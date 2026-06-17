# hri_full_v1.py — HR-ассистент с динамическими подразделениями
# Запуск: streamlit run hri_full_v1.py

import streamlit as st
import os
import yaml
import json
import time
import requests
import boto3
import subprocess
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook
import PyPDF2
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import mm
import tempfile
from datetime import datetime

# ============================================================
# ОСНОВНЫЕ КОНСТАНТЫ
# ============================================================
CLIENTS = {0: "Админ"}  # Будет дополняться из departments.json

# ============================================================
# УПРАВЛЕНИЕ ПОДРАЗДЕЛЕНИЯМИ
# ============================================================
DEPARTMENTS_FILE = "data/departments.json"

def load_departments():
    """Загружает список подразделений из файла"""
    if not os.path.exists(DEPARTMENTS_FILE):
        # Создаём с дефолтными подразделениями
        default = {
            "departments": [
                {"id": 1, "name": "Маркетинг", "slug": "marketing"},
                {"id": 2, "name": "Продажи", "slug": "sales"},
                {"id": 3, "name": "Бухгалтерия", "slug": "accounting"},
                {"id": 4, "name": "Склад", "slug": "warehouse"},
                {"id": 5, "name": "Логистика", "slug": "logistics"},
                {"id": 6, "name": "Руководители", "slug": "managers"},
                {"id": 99, "name": "Тестовый", "slug": "test"}
            ]
        }
        save_departments(default)
        return default["departments"]
    with open(DEPARTMENTS_FILE, "r", encoding="utf-8") as f:
        data = json.load(f)
        return data.get("departments", [])

def save_departments(departments_data):
    """Сохраняет список подразделений в файл"""
    os.makedirs("data", exist_ok=True)
    with open(DEPARTMENTS_FILE, "w", encoding="utf-8") as f:
        json.dump(departments_data, f, ensure_ascii=False, indent=2)

def add_department(name):
    """Добавляет новое подразделение и возвращает его ID"""
    departments = load_departments()
    new_id = max([d["id"] for d in departments], default=0) + 1
    slug = name.lower().replace(" ", "_")
    departments.append({"id": new_id, "name": name, "slug": slug})
    save_departments({"departments": departments})
    return new_id

def get_department_by_id(dept_id):
    """Возвращает подразделение по ID"""
    departments = load_departments()
    for d in departments:
        if d["id"] == dept_id:
            return d
    return None

# ============================================================
# НАСТРОЙКИ ТЕМЫ И КАСТОМИЗАЦИИ ИНТЕРФЕЙСА
# ============================================================
if "theme" not in st.session_state:
    st.session_state.theme = "light"
if "accent_color" not in st.session_state:
    st.session_state.accent_color = "#00a3ff"
if "compact_mode" not in st.session_state:
    st.session_state.compact_mode = False

def apply_theme_and_css():
    """Применяет выбранную тему, цвет акцента и компактный режим через CSS."""
    if st.session_state.theme == "dark":
        bg = "#0e1111"
        text = "#ffffff"
        card_bg = "#262730"
        secondary_bg = "#1a1c23"
    elif st.session_state.theme == "system":
        bg = "var(--st-bg)"
        text = "var(--st-text)"
        card_bg = "var(--st-card-bg)"
        secondary_bg = "var(--st-secondary-bg)"
    else:  # light
        bg = "#ffffff"
        text = "#0e1111"
        card_bg = "#f0f2f6"
        secondary_bg = "#e6e9ef"

    accent = st.session_state.accent_color
    compact_css = """
        div.block-container {padding-top: 0.5rem; padding-bottom: 0.5rem;}
        .stMarkdown p, .stTextArea label, .stSelectbox label {margin-bottom: 0.2rem;}
        .stButton button {padding: 0.3rem 0.8rem;}
    """ if st.session_state.compact_mode else ""

    st.markdown(f"""
    <style>
        .stApp {{ background-color: {bg}; }}
        .reportview-container .main .block-container {{
            background-color: {bg};
            color: {text};
        }}
        .st-bw, .st-cb, .st-at {{ background-color: {card_bg} !important; }}
        .stButton button, .stDownloadButton button {{
            background-color: {accent} !important;
            border-color: {accent} !important;
        }}
        .stButton button:hover, .stDownloadButton button:hover {{ filter: brightness(0.9); }}
        .custom-card {{
            background-color: {card_bg};
            border-radius: 10px;
            padding: 1rem;
            margin: 1rem 0;
            border-left: 4px solid {accent};
        }}
        {compact_css}
    </style>
    """, unsafe_allow_html=True)

#apply_theme_and_css()

# ============================================================
# ЗАГРУЗКА КОНФИГУРАЦИИ И КЛЮЧЕЙ
# ============================================================
load_dotenv()

def load_config(config_path="hri_full_v1_config.yaml"):
    with open(config_path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)

config = load_config()

api_key_env = config["model"].get("api_key_env", "PROXYAPI_KEY")
api_key = os.getenv(api_key_env)
base_url = config["model"].get("base_url", "https://api.proxyapi.ru/openai/v1")

client = OpenAI(
    base_url=base_url,
    api_key=api_key
)

# ============================================================
# ВСТРОЕННЫЕ ОБРАЗЦЫ ДОКУМЕНТОВ
# ============================================================
EXAMPLE_PROFILE = """
Профиль должности (указать должность)
Описание
Комментарии
Общие данные
Подразделение

Непосредственный руководитель сотрудника

ЦКП (ценный конечный продукт - за что именно платятся деньги) должности

Основные обязанности

Анкетные данные
Предыдущие места работы для нужного опыта

Предыдущие должности для нужного опыта

Образование: 

Пол:  

Возраст:  

Семейное положение, дети:

Место жительства:

Прочие требования: 

Стоп факторы: 
Например, очное обучение; неготовность к офисному формату

Hard Skills, без которых кандидат не сможет работать:
Навык, умение или знание: 
Как проявляется на стадии собеседований (что хотим услышать):

Hard Skills, которые толковый кандидат может освоить в первые 1-6 недель работы:

Личные качества
Качество/Как проявляется
Как увидим на стадии собеседований?

Условия работы:
Формат, режим
Оклад (несгораемая цифра)
Переменная часть - бонус
Длительность испытательного срока
Зарплатные условия на период испытательного срока
"""

EXAMPLE_VACANCY = """
Младший менеджер по маркетплейсам (Wildberries)
от 75 000 до 85 000 ₽ за месяц, до вычета налогов
Торговый дом
Ладожская, Санкт-Петербург, проспект Энергетиков, 4к1

Выплаты: два раза в месяц
Опыт работы: не требуется
Оформление: Трудовой договор
График: 5/2
Формат работы: на месте работодателя

YourBox — крупный бренд обуви из Санкт-Петербурга с 10-летней историей.

Основные задачи:
— Создание и редактирование карточек товаров на платформе;
— Проведение простых А/Б тестов для повышения конверсии;
— Работа с отзывами и баллами за отзывы;
— Создание баннерных рекламных кампаний и управление ими (WB.Media);
— Организация фотосессий товаров;
— Работа с фото- и видеоконтентом;
— SEO-оптимизация карточек товаров.

Требования к кандидату:
— Опыт работы с карточками товаров на Wildberries;
— Внимательность к деталям;
— Ответственный подход к задачам.

Мы предлагаем:
— современное оборудованное пространство;
— обучение и поддержку;
— возможности для роста;
— возможность приобрести обувь по себестоимости.
"""

EXAMPLE_QUESTIONNAIRE = """
№   Вопрос                                                                                         Желательный результат
1    Почему вы сейчас ищете новую работу?                                                         Объективная причина или стремление к росту.
2    Что для вас важно в работе, и какие условия вы ищете?                                        Интерес к задачам компании, без излишнего фокуса на комфорт.
3    Какой у вас опыт работы с карточками товаров на Wildberries?                                 Опыт от 1 года, работа с 30+ артикулами.
4    Опишите ситуацию, когда вы заметили ошибку в карточке товара и как её исправили.            Конкретный пример, внимательность.
5    Как вы оцениваете эффективность рекламной кампании на Wildberries?                          Знание метрик CTR, CR, ACoS.
6    Как вы организуете свою работу при ведении большого количества товаров?                     Приоритизация, планирование.
7    Где вы учились работе с маркетплейсами?                                                      Курсы, самообучение, практика.
8    Итог:                                                                                       Проходной балл.
"""

# ============================================================
# ФУНКЦИИ ДЛЯ РАБОТЫ С ФАЙЛАМИ И ТЕКСТОМ
# ============================================================
def extract_text(uploaded_file):
    filename = uploaded_file.name.lower()
    if filename.endswith(".txt"):
        return uploaded_file.read().decode("utf-8")
    elif filename.endswith(".docx"):
        doc = Document(uploaded_file)
        return "\n".join([p.text for p in doc.paragraphs])
    elif filename.endswith(".xlsx"):
        wb = load_workbook(uploaded_file)
        sheet = wb.active
        rows = []
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
            rows.append(row_text)
        return "\n".join(rows)
    else:
        raise ValueError("Неподдерживаемый формат файла")

def get_direct_yandex_link(public_url):
    if "disk.yandex.ru" in public_url:
        return public_url.replace("/i/", "/d/")
    elif "yadi.sk" in public_url:
        return public_url.replace("/i/", "/d/")
    return public_url

def extract_text_from_pdf_url(pdf_url):
    try:
        download_url = get_direct_yandex_link(pdf_url)
        response = requests.get(download_url, timeout=30)
        if response.status_code == 200:
            pdf_file = BytesIO(response.content)
            reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        return None
    except Exception as e:
        print(f"Ошибка извлечения текста из PDF: {e}")
        return None

# ============================================================
# ФУНКЦИИ ДЛЯ ЯНДЕКС ОБЛАКА
# ============================================================
def upload_to_s3_and_get_url(local_path, bucket, access_key, secret_key):
    session = boto3.session.Session()
    s3 = session.client(
        service_name='s3',
        endpoint_url='https://storage.yandexcloud.net',
        aws_access_key_id=access_key,
        aws_secret_access_key=secret_key
    )
    object_name = os.path.basename(local_path)
    file_size = os.path.getsize(local_path)
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    def callback(bytes_transferred):
        progress = int(bytes_transferred / file_size * 100)
        progress_bar.progress(progress)
        status_text.text(f"Загрузка в облако: {progress}%")
    
    s3.upload_file(local_path, bucket, object_name, Callback=callback)
    progress_bar.empty()
    status_text.empty()
    
    return f"https://storage.yandexcloud.net/{bucket}/{object_name}"

def convert_to_pcm(input_path, output_path=None):
    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + "_speechkit.pcm"
    subprocess.run([
        "ffmpeg", "-y", "-i", input_path,
        "-acodec", "pcm_s16le", "-ac", "1", "-ar", "16000",
        "-f", "s16le", output_path
    ], check=True, capture_output=True)
    return output_path

def recognize_long_audio(audio_url, api_key):
    response = requests.post(
        "https://transcribe.api.cloud.yandex.net/speech/stt/v2/longRunningRecognize",
        headers={"Authorization": f"Api-Key {api_key}"},
        json={
            "config": {
                "specification": {
                    "languageCode": "ru-RU",
                    "profanityFilter": "false",
                    "audioEncoding": "LINEAR16_PCM",
                    "sampleRateHertz": 16000,
                    "audioChannelCount": 1
                }
            },
            "audio": {"uri": audio_url}
        }
    )
    if response.status_code != 200:
        raise Exception(f"Ошибка запроса: {response.text}")
    operation = response.json()
    operation_id = operation["id"]
    while True:
        time.sleep(5)
        status_response = requests.get(
            f"https://operation.api.cloud.yandex.net/operations/{operation_id}",
            headers={"Authorization": f"Api-Key {api_key}"}
        )
        status_data = status_response.json()
        if status_data.get("done"):
            if "error" in status_data:
                raise Exception(f"Ошибка распознавания: {status_data['error']}")
            full_text = ""
            for chunk in status_data["response"]["chunks"]:
                full_text += chunk["alternatives"][0]["text"] + " "
            return full_text.strip()

# ============================================================
# УПРАВЛЕНИЕ ДАННЫМИ: ВАКАНСИИ, ЧАТЫ, ИСТОРИЯ
# ============================================================
VACANCIES_FILE = "data/vacancies_db.json"
CHATS_FILE = "data/chats_db.json"
HISTORY_DIR = "data/history"
HISTORY_INDEX_FILE = os.path.join(HISTORY_DIR, "index.json")

def ensure_history_dir():
    os.makedirs(HISTORY_DIR, exist_ok=True)

def save_generation_to_history(generated_data, transcript_text=None, vacancy_title=None):
    ensure_history_dir()
    timestamp = int(time.time())
    dt_str = time.strftime("%Y%m%d_%H%M%S", time.localtime(timestamp))
    title = generated_data.get("должность", "unknown").replace(" ", "_")
    if vacancy_title:
        title = vacancy_title.replace(" ", "_")
    filename = f"{dt_str}_{title}.json"
    filepath = os.path.join(HISTORY_DIR, filename)
    
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(generated_data, f, ensure_ascii=False, indent=2)
    
    if os.path.exists(HISTORY_INDEX_FILE):
        with open(HISTORY_INDEX_FILE, "r", encoding="utf-8") as f:
            index = json.load(f)
    else:
        index = []
    
    record = {
        "timestamp": timestamp,
        "datetime": dt_str,
        "filename": filename,
        "title": generated_data.get("должность", ""),
        "preview": generated_data.get("текст_вакансии", "")[:100] + "..." if generated_data.get("текст_вакансии") else "",
        "vacancy_title": vacancy_title or "",
        "has_transcript": bool(transcript_text)
    }
    index.append(record)
    index.sort(key=lambda x: x["timestamp"], reverse=True)
    with open(HISTORY_INDEX_FILE, "w", encoding="utf-8") as f:
        json.dump(index, f, ensure_ascii=False, indent=2)
    return record

def load_generation_from_history(filename):
    filepath = os.path.join(HISTORY_DIR, filename)
    if os.path.exists(filepath):
        with open(filepath, "r", encoding="utf-8") as f:
            return json.load(f)
    return None

def delete_generation_from_history(filename):
    filepath = os.path.join(HISTORY_DIR, filename)
    if os.path.exists(filepath):
        os.remove(filepath)
    if os.path.exists(HISTORY_INDEX_FILE):
        with open(HISTORY_INDEX_FILE, "r", encoding="utf-8") as f:
            index = json.load(f)
        new_index = [rec for rec in index if rec["filename"] != filename]
        with open(HISTORY_INDEX_FILE, "w", encoding="utf-8") as f:
            json.dump(new_index, f, ensure_ascii=False, indent=2)
        return True
    return False

def get_history_index():
    if os.path.exists(HISTORY_INDEX_FILE):
        with open(HISTORY_INDEX_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

def load_chats():
    if not os.path.exists(CHATS_FILE):
        return []
    with open(CHATS_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

def save_chats(chats_list):
    os.makedirs("data", exist_ok=True)
    with open(CHATS_FILE, "w", encoding="utf-8") as f:
        json.dump(chats_list, f, ensure_ascii=False, indent=2)

def load_vacancies():
    if not os.path.exists(VACANCIES_FILE):
        return []
    with open(VACANCIES_FILE, "r", encoding="utf-8") as f:
        data = json.load(f)
        vacancies = data.get("vacancies", [])
    migrated = False
    for v in vacancies:
        if "documents" not in v:
            v["documents"] = {
                "profile": "",
                "vacancy_text": "",
                "questions": "",
                "keywords": "",
                "notes": ""
            }
            migrated = True
    if migrated:
        save_vacancies(vacancies)
    return vacancies

def save_vacancies(vacancies_list):
    os.makedirs("data", exist_ok=True)
    with open(VACANCIES_FILE, "w", encoding="utf-8") as f:
        json.dump({"vacancies": vacancies_list}, f, ensure_ascii=False, indent=2)

def create_vacancy(title, chat_id, client_id=0):
    vacancies = load_vacancies()
    if any(v["title"] == title for v in vacancies):
        return False, "Вакансия с таким названием уже существует"
    new_vacancy = {
        "id": len(vacancies) + 1,
        "title": title,
        "chat_id": chat_id,
        "client_id": client_id,
        "documents": {
            "profile": "",
            "vacancy_text": "",
            "questions": "",
            "keywords": "",
            "notes": ""
        },
        "candidates": []
    }
    vacancies.append(new_vacancy)
    save_vacancies(vacancies)
    return True, new_vacancy

def update_vacancy_docs(vacancy_title, docs_dict):
    vacancies = load_vacancies()
    for v in vacancies:
        if v["title"] == vacancy_title:
            if "documents" not in v:
                v["documents"] = {
                    "profile": "", "vacancy_text": "", "questions": "", "keywords": "", "notes": ""
                }
            v["documents"].update(docs_dict)
            save_vacancies(vacancies)
            return True
    return False

def delete_vacancy(vacancy_title):
    vacancies = load_vacancies()
    new_vacancies = [v for v in vacancies if v["title"] != vacancy_title]
    if len(new_vacancies) < len(vacancies):
        save_vacancies(new_vacancies)
        return True
    return False

# ============================================================
# ФУНКЦИИ TELEGRAM
# ============================================================
def send_telegram_message(bot_token, chat_id, text):
    url = f"https://api.telegram.org/bot{bot_token}/sendMessage"
    payload = {"chat_id": chat_id, "text": text, "parse_mode": "HTML"}
    try:
        response = requests.post(url, json=payload)
        data = response.json()
        if data.get("ok"):
            return True, "Уведомление доставлено в Telegram"
        else:
            return False, f"Ошибка Telegram: {data.get('description', 'неизвестно')}"
    except Exception as e:
        return False, f"Сетевая ошибка: {e}"

# ============================================================
# ОЦЕНКА КАНДИДАТА ИИ
# ============================================================
def evaluate_candidate_with_ai(resume_text, transcript_text, job_title):
    prompt = f"""
Ты — опытный HR-директор. Оцени кандидата на позицию: {job_title}

РЕЗЮМЕ:
{resume_text}

РАСШИФРОВКА:
{transcript_text if transcript_text else "Не предоставлена"}

Верни JSON: {{"score": 7, "comment": "...", "strengths": [], "weaknesses": []}}
"""
    try:
        response = client.chat.completions.create(
            model=config["model"]["name"],
            messages=[
                {"role": "system", "content": "Ты HR-эксперт. Отвечай строго в JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=config["model"]["max_tokens"]
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        return {"score": 0, "comment": f"Ошибка: {e}", "strengths": [], "weaknesses": []}

# ============================================================
# НОРМАЛИЗАЦИЯ JSON
# ============================================================
def normalize_docs(doc):
    if not isinstance(doc, dict):
        return doc
    if 'профиль' in doc and isinstance(doc['профиль'], dict):
        profile = doc['профиль']
        for field in ['обязательные_требования', 'желательные_требования', 'психологические_черты']:
            if field in profile and isinstance(profile[field], list):
                new_items = []
                for item in profile[field]:
                    if isinstance(item, str):
                        if field == 'психологические_черты':
                            new_items.append({"качество": item, "проявление": item})
                        else:
                            new_items.append({"навык": item, "описание": item})
                    elif isinstance(item, dict):
                        new_items.append(item)
                    else:
                        if field == 'психологические_черты':
                            new_items.append({"качество": str(item), "проявление": str(item)})
                        else:
                            new_items.append({"навык": str(item), "описание": str(item)})
                profile[field] = new_items
        if 'задачи' in profile and isinstance(profile['задачи'], list):
            profile['задачи'] = [str(t) for t in profile['задачи']]
    if 'опросник' in doc and isinstance(doc['опросник'], list):
        new_q = []
        for q in doc['опросник']:
            if isinstance(q, str):
                new_q.append({"вопрос": q, "пример_ответа": ""})
            elif isinstance(q, dict):
                new_q.append({
                    "вопрос": q.get('вопрос', q.get('question', str(q))),
                    "пример_ответа": q.get('пример_ответа', q.get('example', ''))
                })
            else:
                new_q.append({"вопрос": str(q), "пример_ответа": ""})
        doc['опросник'] = new_q
    if 'ключевые_слова' in doc and isinstance(doc['ключевые_слова'], list):
        doc['ключевые_слова'] = [str(kw) for kw in doc['ключевые_слова']]
    return doc

# ============================================================
# ЭКСПОРТ В WORD И PDF
# ============================================================
def export_to_word(gen_data):
    doc = Document()
    title = doc.add_heading(f"{config['app']['name']} v{config['app']['version']}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading("ПРОФИЛЬ ДОЛЖНОСТИ", level=1)
    profile = gen_data.get("профиль", {})
    doc.add_heading(f"Должность: {gen_data.get('должность', '—')}", level=2)
    doc.add_paragraph(f"Подразделение: {profile.get('подразделение', '—')}")
    doc.add_paragraph(f"Непосредственный руководитель: {profile.get('непосредственный_руководитель', '—')}")
    
    doc.add_heading("1. Задачи", level=2)
    tasks = profile.get("задачи", [])
    for task in tasks:
        doc.add_paragraph(task, style='List Bullet')
    
    doc.add_heading("2. Анкетные требования", level=2)
    at = profile.get("анкетные_требования", {})
    doc.add_paragraph(f"Возраст: {at.get('возраст', '—')}")
    doc.add_paragraph(f"Пол: {at.get('пол', '—')}")
    doc.add_paragraph("Стоп-факторы:")
    for sf in at.get("стоп_факторы", []):
        doc.add_paragraph(sf, style='List Bullet')
    
    doc.add_heading("3. Обязательные требования (Hard Skills)", level=2)
    for req in profile.get("обязательные_требования", []):
        if isinstance(req, dict):
            doc.add_paragraph(f"{req.get('навык', '')}: {req.get('описание', '')}", style='List Bullet')
        else:
            doc.add_paragraph(str(req), style='List Bullet')
    
    doc.add_heading("4. Желательные требования", level=2)
    for req in profile.get("желательные_требования", []):
        if isinstance(req, dict):
            doc.add_paragraph(f"{req.get('навык', '')}: {req.get('описание', '')}", style='List Bullet')
        else:
            doc.add_paragraph(str(req), style='List Bullet')
    
    doc.add_heading("5. Психологические черты (Soft Skills)", level=2)
    for trait in profile.get("психологические_черты", []):
        if isinstance(trait, dict):
            doc.add_paragraph(f"{trait.get('качество', '')}: {trait.get('проявление', '')}", style='List Bullet')
        else:
            doc.add_paragraph(str(trait), style='List Bullet')
    
    doc.add_heading("6. Условия работы", level=2)
    cond = profile.get("условия_работы", {})
    doc.add_paragraph(f"Формат: {cond.get('формат', '—')}")
    doc.add_paragraph(f"Режим: {cond.get('режим', '—')}")
    doc.add_paragraph(f"Зарплата: {cond.get('зарплата', '—')}")
    doc.add_paragraph(f"Испытательный срок: {cond.get('испытательный_срок', '—')}")
    
    doc.add_page_break()
    doc.add_heading("ТЕКСТ ВАКАНСИИ", level=1)
    doc.add_paragraph(gen_data.get("текст_вакансии", ""))
    
    doc.add_page_break()
    doc.add_heading("ОПРОСНИК ДЛЯ ПЕРВИЧНОГО СОБЕСЕДОВАНИЯ", level=1)
    for i, q in enumerate(gen_data.get("опросник", []), 1):
        doc.add_heading(f"{i}. {q.get('вопрос', '')}", level=2)
        doc.add_paragraph(f"Пример желательного ответа: {q.get('пример_ответа', '')}", style='Intense Quote')
    
    doc.add_heading("КЛЮЧЕВЫЕ СЛОВА ДЛЯ ПОИСКА", level=1)
    keywords = gen_data.get("ключевые_слова", [])
    doc.add_paragraph(", ".join(keywords))
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        return tmp.name

def export_to_pdf(gen_data):
    font_dir = "fonts"
    os.makedirs(font_dir, exist_ok=True)
    font_path = os.path.join(font_dir, "DejaVuSans.ttf")
    
    if not os.path.exists(font_path):
        st.error(
            "❌ Не найден шрифт для генерации PDF.\n\n"
            "Скачайте файл DejaVuSans.ttf и поместите его в папку 'fonts' в корне проекта."
        )
        return None

    pdfmetrics.registerFont(TTFont('DejaVu', font_path))
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        pdf_path = tmp.name

    c = canvas.Canvas(pdf_path, pagesize=A4)
    width, height = A4
    y = height - 20
    line_height = 14

    def draw_text(text, font_size=11, bold=False):
        nonlocal y
        if y < 50:
            c.showPage()
            y = height - 20
        c.setFont('DejaVu', font_size)
        lines = text.split('\n')
        for line in lines:
            if y < 50:
                c.showPage()
                y = height - 20
            c.drawString(20, y, line)
            y -= line_height

    def draw_bullet_list(items):
        nonlocal y
        for item in items:
            if y < 50:
                c.showPage()
                y = height - 20
            c.setFont('DejaVu', 11)
            c.drawString(25, y, f"• {item}")
            y -= line_height

    def draw_section_title(title):
        nonlocal y
        if y < 60:
            c.showPage()
            y = height - 20
        c.setFont('DejaVu', 14)
        c.drawString(20, y, title)
        y -= line_height + 5

    draw_section_title("ПРОФИЛЬ ДОЛЖНОСТИ")
    profile = gen_data.get("профиль", {})
    draw_text(f"Должность: {gen_data.get('должность', '—')}", font_size=12)
    draw_text(f"Подразделение: {profile.get('подразделение', '—')}")
    draw_text(f"Руководитель: {profile.get('непосредственный_руководитель', '—')}")

    draw_section_title("1. Задачи")
    draw_bullet_list(profile.get("задачи", []))

    draw_section_title("2. Анкетные требования")
    at = profile.get("анкетные_требования", {})
    draw_text(f"Возраст: {at.get('возраст', '—')}")
    draw_text(f"Пол: {at.get('пол', '—')}")
    draw_text("Стоп-факторы:")
    draw_bullet_list(at.get("стоп_факторы", []))

    draw_section_title("3. Обязательные требования")
    for req in profile.get("обязательные_требования", []):
        if isinstance(req, dict):
            draw_text(f"• {req.get('навык', '')}: {req.get('описание', '')}")
        else:
            draw_text(f"• {req}")

    draw_section_title("4. Желательные требования")
    for req in profile.get("желательные_требования", []):
        if isinstance(req, dict):
            draw_text(f"• {req.get('навык', '')}: {req.get('описание', '')}")
        else:
            draw_text(f"• {req}")

    draw_section_title("5. Психологические черты")
    for trait in profile.get("психологические_черты", []):
        if isinstance(trait, dict):
            draw_text(f"• {trait.get('качество', '')}: {trait.get('проявление', '')}")
        else:
            draw_text(f"• {trait}")

    draw_section_title("6. Условия работы")
    cond = profile.get("условия_работы", {})
    draw_text(f"Формат: {cond.get('формат', '—')}")
    draw_text(f"Режим: {cond.get('режим', '—')}")
    draw_text(f"Зарплата: {cond.get('зарплата', '—')}")
    draw_text(f"Испытательный срок: {cond.get('испытательный_срок', '—')}")

    draw_section_title("ТЕКСТ ВАКАНСИИ")
    draw_text(gen_data.get("текст_вакансии", ""))

    draw_section_title("ОПРОСНИК ДЛЯ СОБЕСЕДОВАНИЯ")
    for i, q in enumerate(gen_data.get("опросник", []), 1):
        draw_text(f"{i}. {q.get('вопрос', '')}", font_size=11)
        draw_text(f"   Пример: {q.get('пример_ответа', '')}", font_size=10)
        y -= 5

    draw_section_title("КЛЮЧЕВЫЕ СЛОВА")
    draw_text(", ".join(gen_data.get("ключевые_слова", [])))

    c.save()
    return pdf_path

# ============================================================
# ИНТЕРФЕЙС STREAMLIT
# ============================================================
st.set_page_config(
    page_title=f"{config['app']['name']} v{config['app']['version']}",
    page_icon="🧠",
    layout="wide",
)

st.title(f"🧠 {config['app']['name']} v{config['app']['version']}")
st.markdown("**Разработчик:** А.А. Крупин")
st.caption("Загрузите любой файл с информацией о вакансии (аудио-, видео-, текст) и получите профессиональный профиль должности, текст вакансии, опросник и ключевые слова.")

# Боковая панель
with st.sidebar:
    st.header("⚙️ Конфигурация")
    st.write(f"**Модель:** `{config['model']['name']}`")
    st.write(f"**Температура:** {config['model']['temperature']}")

    st.divider()
    if st.button("🔄 Перезагрузить конфиг"):
        config = load_config()
        st.rerun()

    # Динамическая навигация по подразделениям
    st.sidebar.divider()
st.sidebar.subheader("👥 Клиентские зоны")
departments = load_departments()
for dept in departments:
    # Создаём прямую ссылку на страницу клиента с параметром dept
    #st.sidebar.markdown(f"🏢 [{dept['name']}](/client?dept={dept['name']})")
    # Альтернативно, можно использовать link_button (если ваша версия Streamlit поддерживает)
    st.sidebar.link_button(f"🏢 {dept['name']}", f"/client?dept={dept['name']}")

# Вкладки
tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(["🎤 Расшифровка", "📝 Генерация", "📄 Результаты", "🎯 Воронка кандидатов", "📖 Инструкции", "📜 История"])

# ---------- ВКЛАДКА 1: РАСШИФРОВКА ----------
with tab1:
    st.header("📥 Получение текста")
    source = st.radio("Выберите источник текста:", ("Аудио/видео", "Готовый файл"), horizontal=True, key="source_radio")

    if source == "Аудио/видео":
        method = st.radio("Метод расшифровки:", ("Яндекс (SpeechKit)",), horizontal=True, key="method_radio")
        uploaded_audio = st.file_uploader("Выберите аудио/видео", type=["mp3", "mp4", "wav", "webm", "mkv", "ogg"])
        if uploaded_audio:
            os.makedirs("data/tmp", exist_ok=True)
            audio_path = os.path.join("data/tmp", uploaded_audio.name)
            with open(audio_path, "wb") as f:
                f.write(uploaded_audio.read())
            if st.button("🎙️ Расшифровать"):
                with st.spinner("Расшифровка..."):
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    for i in range(100):
                        time.sleep(0.05)
                        progress_bar.progress(i + 1)
                        status_text.text(f"Расшифровка... {i+1}%")
                    if method == "Яндекс (SpeechKit)":
                        pass
                        result = model.transcribe(audio_path, language="ru", task="transcribe", fp16=False)
                        transcript_text = result["text"]
                    else:
                        try:
                            pcm_path = convert_to_pcm(audio_path)
                            audio_url = upload_to_s3_and_get_url(pcm_path,
                                os.getenv("YANDEX_BUCKET_NAME"),
                                os.getenv("YANDEX_ACCESS_KEY_ID"),
                                os.getenv("YANDEX_SECRET_ACCESS_KEY"))
                            transcript_text = recognize_long_audio(audio_url, os.getenv("YANDEX_API_KEY"))
                        except Exception as e:
                            st.error(f"Ошибка: {e}")
                            transcript_text = ""
                    if transcript_text:
                        st.session_state.transcript = transcript_text
                        base = os.path.splitext(uploaded_audio.name)[0]
                        output_path = os.path.join(config["paths"]["output_dir"], f"{base}_transcript.txt")
                        os.makedirs(config["paths"]["output_dir"], exist_ok=True)
                        with open(output_path, "w", encoding="utf-8") as f:
                            f.write(transcript_text)
                        st.success("Расшифровка завершена!")
                        st.rerun()
    else:
        uploaded_text_file = st.file_uploader("Загрузите файл с текстом", type=["txt", "docx", "xlsx"], key="text_uploader")
        if uploaded_text_file:
            try:
                text_content = extract_text(uploaded_text_file)
                if text_content.strip():
                    st.session_state.transcript = text_content
                    st.success("Файл загружен! Теперь вы можете отредактировать текст ниже.")
                else:
                    st.warning("Файл пуст.")
            except Exception as e:
                st.error(f"Ошибка: {e}")

    if "transcript" in st.session_state and st.session_state.transcript:
        edited_text = st.text_area("Текст для генерации (можно редактировать)", st.session_state.transcript, height=300)
        if st.button("💾 Зафиксировать правки"):
            st.session_state.transcript = edited_text
            st.success("Текст обновлён!")
    else:
        if source == "Аудио/видео":
            st.info("Загрузите аудио/видео и нажмите «Расшифровать».")
        else:
            st.info("Загрузите файл .txt, .docx или .xlsx.")

# ---------- ВКЛАДКА 2: ГЕНЕРАЦИЯ ----------
with tab2:
    st.header("📝 Генерация документов по шаблонам")
    if "transcript" not in st.session_state or not st.session_state.transcript.strip():
        st.warning("Сначала получите текст в первой вкладке.")
    else:
        if st.button("✨ Сгенерировать всё"):
            progress_bar = st.progress(0)
            status_text = st.empty()
            for i in range(100):
                time.sleep(0.02)
                progress_bar.progress(i + 1)
                status_text.text(f"Генерация документов... {i+1}%")
            progress_bar.empty()
            status_text.empty()

            with st.spinner("ИИ создаёт документы..."):
                system_prompt = f"""
Ты — HR-директор с 15-летним опытом работы в подборе кандидатов в e-commerce сфере. 
На основе текстовой расшифровки разговора с заказчиком создай:
- профиль должности,
- текст вакансии,
- вопросы для первичного собеседования (опросник),
- ключевые слова для поиска кандидатов на hh.ru и в соцсетях.

Ниже приведены образцы документов. Строго соблюдай стиль и структуру образцов.

Образец профиля:
{EXAMPLE_PROFILE}

Образец текста вакансии:
{EXAMPLE_VACANCY}

Образец опросника:
{EXAMPLE_QUESTIONNAIRE}

**Требования к опроснику:**
- Обязательно не менее 7 вопросов.
- Первые два вопроса: "Почему ищете работу?" и "Что для вас важно в работе, и какие условия вы ищете?".
- Для каждого вопроса дай реалистичный пример желательного ответа.

**Формат ответа:**
Верни строго JSON без лишних пояснений.
"""
                user_message = f"Текст расшифровки:\n{st.session_state.transcript}"

                try:
                    response = client.chat.completions.create(
                        model=config["model"]["name"],
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_message}
                        ],
                        temperature=config["model"]["temperature"],
                        max_tokens=config["model"]["max_tokens"]
                    )
                    content = response.choices[0].message.content
                    if "```json" in content:
                        content = content.split("```json")[1].split("```")[0]
                    elif "```" in content:
                        content = content.split("```")[1].split("```")[0]
                    result = json.loads(content.strip())
                    result = normalize_docs(result)
                    st.session_state.generated = result
                    timestamp = str(int(time.time()))
                    out_json = os.path.join(config["paths"]["output_dir"], f"generated_{timestamp}.json")
                    os.makedirs(config["paths"]["output_dir"], exist_ok=True)
                    with open(out_json, "w", encoding="utf-8") as f:
                        json.dump(result, f, ensure_ascii=False, indent=2)
                    save_generation_to_history(result, st.session_state.transcript)
                    st.success("Документы созданы!")
                    st.rerun()
                except Exception as e:
                    st.error(f"Ошибка: {e}")

        if "generated" in st.session_state:
            st.divider()
            st.subheader("✏️ Корректировка через ИИ")
            correction = st.text_area("Дополнительные указания:", height=80)
            if st.button("🔄 Доработать"):
                if correction.strip():
                    with st.spinner("Доработка..."):
                        current_json = json.dumps(st.session_state.generated, ensure_ascii=False)
                        refine_msg = f"Учти: {correction}\n\nТекущие документы:\n{current_json}\nВерни полный JSON."
                        try:
                            response = client.chat.completions.create(
                                model=config["model"]["name"],
                                messages=[
                                    {"role": "system", "content": "Ты HR-директор. Обнови JSON по запросу."},
                                    {"role": "user", "content": refine_msg}
                                ],
                                temperature=config["model"]["temperature"],
                                max_tokens=config["model"]["max_tokens"]
                            )
                            content = response.choices[0].message.content
                            if "```json" in content:
                                content = content.split("```json")[1].split("```")[0]
                            elif "```" in content:
                                content = content.split("```")[1].split("```")[0]
                            new_result = json.loads(content.strip())
                            new_result = normalize_docs(new_result)
                            st.session_state.generated = new_result
                            st.success("Обновлено!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Ошибка: {e}")

            st.divider()
            st.subheader("💾 Сохранить в вакансию")
            vacancies = load_vacancies()
            if vacancies:
                target = st.selectbox("Выберите существующую вакансию", [v["title"] for v in vacancies])
                if st.button("Обновить документы в выбранной вакансии"):
                    docs = st.session_state.generated
                    profile_text = json.dumps(docs.get("профиль", {}), ensure_ascii=False, indent=2)
                    vacancy_text = docs.get("текст_вакансии", "")
                    questions_text = json.dumps(docs.get("опросник", []), ensure_ascii=False, indent=2)
                    keywords_text = ", ".join(docs.get("ключевые_слова", []))
                    if update_vacancy_docs(target, {
                        "profile": profile_text,
                        "vacancy_text": vacancy_text,
                        "questions": questions_text,
                        "keywords": keywords_text
                    }):
                        save_generation_to_history(st.session_state.generated, vacancy_title=target)
                        st.success(f"Сохранено в «{target}»!")
                        st.rerun()
                    else:
                        st.error("Вакансия не найдена")
            else:
                st.info("Нет вакансий. Создайте новую ниже.")

            with st.expander("➕ Создать новую вакансию с этими документами"):
                chats = load_chats()
                if chats:
                    chat_opts = {c["name"]: c["id"] for c in chats}
                    selected_chat_name = st.selectbox("Чат Telegram", list(chat_opts.keys()))
                    chat_id = chat_opts[selected_chat_name]
                    # Получаем client_id из настроек чата
                    selected_chat = next((c for c in chats if c["name"] == selected_chat_name), None)
                    client_id = selected_chat.get("department_id", 0) if selected_chat else 0
                    dept_name = selected_chat.get("department_name", "Админ") if selected_chat else "Админ"
                    st.info(f"👥 Подразделение: {dept_name}")
                else:
                    chat_id = ""
                    client_id = 0
                    st.warning("Нет чатов. Добавьте в четвёртой вкладке.")
                new_title = st.text_input("Название должности")
                if st.button("Создать и сохранить"):
                    if new_title.strip() and chat_id:
                        success, _ = create_vacancy(new_title.strip(), chat_id, client_id)
                        if success:
                            docs = st.session_state.generated
                            profile_text = json.dumps(docs.get("профиль", {}), ensure_ascii=False, indent=2)
                            vacancy_text = docs.get("текст_вакансии", "")
                            questions_text = json.dumps(docs.get("опросник", []), ensure_ascii=False, indent=2)
                            keywords_text = ", ".join(docs.get("ключевые_слова", []))
                            update_vacancy_docs(new_title.strip(), {
                                "profile": profile_text,
                                "vacancy_text": vacancy_text,
                                "questions": questions_text,
                                "keywords": keywords_text
                            })
                            save_generation_to_history(st.session_state.generated, vacancy_title=new_title.strip())
                            st.success("Вакансия создана и документы сохранены!")
                            st.rerun()
                        else:
                            st.error("Такое название уже существует.")
                    else:
                        st.warning("Заполните все поля.")

# ---------- ВКЛАДКА 3: РЕЗУЛЬТАТЫ ----------
with tab3:
    st.header("📄 Готовые документы")
    if "generated" in st.session_state:
        gen = st.session_state.generated
        subtab1, subtab2, subtab3, subtab4 = st.tabs(["📌 Профиль", "📝 Вакансия", "📋 Опросник", "🔑 Ключи"])
        with subtab1:
            profile = gen.get("профиль", {})
            st.subheader(f"Должность: {gen.get('должность', '—')}")
            st.markdown(f"**Подразделение:** {profile.get('подразделение', '—') if isinstance(profile, dict) else '—'}")
            st.markdown(f"**Руководитель:** {profile.get('непосредственный_руководитель', '—') if isinstance(profile, dict) else '—'}")
            st.markdown("**Задачи:**")
            tasks = profile.get("задачи", []) if isinstance(profile, dict) else []
            for task in tasks:
                st.markdown(f"- {task}")
            st.markdown("**Анкетные требования:**")
            at = profile.get("анкетные_требования", {}) if isinstance(profile, dict) else {}
            if isinstance(at, dict):
                st.markdown(f"- Возраст: {at.get('возраст', '—')}")
                st.markdown(f"- Пол: {at.get('пол', '—')}")
            st.markdown("**Обязательные требования:**")
            reqs = profile.get("обязательные_требования", []) if isinstance(profile, dict) else []
            for req in reqs:
                if isinstance(req, dict):
                    st.markdown(f"- **{req.get('навык', '')}:** {req.get('описание', '')}")
                else:
                    st.markdown(f"- {req}")
            st.markdown("**Желательные требования:**")
            reqs = profile.get("желательные_требования", []) if isinstance(profile, dict) else []
            for req in reqs:
                if isinstance(req, dict):
                    st.markdown(f"- **{req.get('навык', '')}:** {req.get('описание', '')}")
                else:
                    st.markdown(f"- {req}")
            st.markdown("**Психологические черты:**")
            traits = profile.get("психологические_черты", []) if isinstance(profile, dict) else []
            for trait in traits:
                if isinstance(trait, dict):
                    st.markdown(f"- **{trait.get('качество', '')}:** {trait.get('проявление', '')}")
                else:
                    st.markdown(f"- {trait}")
            st.markdown("**Условия работы:**")
            cond = profile.get("условия_работы", {}) if isinstance(profile, dict) else {}
            if isinstance(cond, dict):
                st.markdown(f"- Формат: {cond.get('формат', '—')}")
                st.markdown(f"- Режим: {cond.get('режим', '—')}")
                st.markdown(f"- Зарплата: {cond.get('зарплата', '—')}")
                st.markdown(f"- Испытательный срок: {cond.get('испытательный_срок', '—')}")
        with subtab2:
            st.text_area("Текст вакансии", gen.get("текст_вакансии", ""), height=300)
        with subtab3:
            questions = gen.get("опросник", [])
            for i, q in enumerate(questions, 1):
                if isinstance(q, dict):
                    st.markdown(f"**{i}. {q.get('вопрос', '')}**")
                    st.caption(f"Пример: {q.get('пример_ответа', '')}")
                else:
                    st.markdown(f"**{i}. {q}**")
        with subtab4:
            keywords = gen.get("ключевые_слова", [])
            st.code(", ".join(keywords) if isinstance(keywords, list) else str(keywords))

        col1, col2, col3 = st.columns(3)
        with col1:
            st.download_button(
                "📥 Скачать JSON",
                data=json.dumps(gen, ensure_ascii=False, indent=2),
                file_name="hr_package.json",
                mime="application/json"
            )
        with col2:
            try:
                word_path = export_to_word(gen)
                with open(word_path, "rb") as f:
                    st.download_button(
                        "📄 Скачать Word",
                        data=f,
                        file_name=f"hr_documents_{gen.get('должность', 'vacancy')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                os.unlink(word_path)
            except Exception as e:
                st.error(f"Ошибка создания Word: {e}")
        with col3:
            try:
                pdf_path = export_to_pdf(gen)
                if pdf_path:
                    with open(pdf_path, "rb") as f:
                        st.download_button(
                            "📑 Скачать PDF",
                            data=f,
                            file_name=f"hr_documents_{gen.get('должность', 'vacancy')}.pdf",
                            mime="application/pdf"
                        )
                    os.unlink(pdf_path)
            except Exception as e:
                st.error(f"Ошибка создания PDF: {e}")
    else:
        st.info("Сгенерируйте документы во второй вкладке.")

# ---------- ВКЛАДКА 4: ВОРОНКА КАНДИДАТОВ ----------
with tab4:
    st.header("🎯 Воронка кандидатов")
    
    # Управление чатами с возможностью создания подразделений
    with st.expander("📂 Мои чаты Telegram"):
        chats = load_chats()
        departments = load_departments()
        
        col1, col2, col3, col4 = st.columns([2, 2, 2, 1])
        with col1:
            new_chat_name = st.text_input("Название чата")
        with col2:
            new_chat_id = st.text_input("Chat ID")
        with col3:
            dept_options = [d["name"] for d in departments]
            dept_choice = st.selectbox(
                "Подразделение",
                dept_options + ["➕ Создать новое..."],
                help="Выберите подразделение или создайте новое"
            )
            new_dept_name = None
            if dept_choice == "➕ Создать новое...":
                new_dept_name = st.text_input("Название нового подразделения")
        
        if st.button("💾 Сохранить чат"):
            if new_chat_name and new_chat_id:
                if dept_choice == "➕ Создать новое...":
                    if new_dept_name:
                        dept_id = add_department(new_dept_name)
                        dept_name = new_dept_name
                    else:
                        st.error("Введите название нового подразделения")
                        st.stop()
                else:
                    dept = next(d for d in departments if d["name"] == dept_choice)
                    dept_id = dept["id"]
                    dept_name = dept["name"]
                
                if not any(c["name"] == new_chat_name for c in chats):
                    chats.append({
                        "name": new_chat_name,
                        "id": new_chat_id,
                        "department_name": dept_name,
                        "department_id": dept_id
                    })
                    save_chats(chats)
                    st.success("Сохранено!")
                    st.rerun()
                else:
                    st.error("Чат с таким именем уже есть.")
        
        if chats:
            st.write("Сохранённые чаты:")
            for i, chat in enumerate(chats):
                col_a, col_b, col_c, col_d = st.columns([2, 2, 2, 1])
                col_a.write(f"**{chat['name']}**")
                col_b.code(chat['id'])
                col_c.write(f"_{chat.get('department_name', '—')}_")
                if col_d.button("🗑️", key=f"del_chat_{i}"):
                    chats.pop(i)
                    save_chats(chats)
                    st.rerun()

    # Создание новой вакансии
    with st.expander("➕ Новая вакансия"):
        new_title = st.text_input("Название должности", key="new_vac4_title")
        chats = load_chats()
        if chats:
            chat_options = {c["name"]: c["id"] for c in chats}
            selected_chat_name = st.selectbox("Чат", list(chat_options.keys()), key="vac4_chat_select")
            chat_id = chat_options[selected_chat_name]
            
            # Находим выбранный чат в списке chats
            selected_chat_obj = None
            for c in chats:
                if c["name"] == selected_chat_name:
                    selected_chat_obj = c
                    break
            
            if selected_chat_obj:
                client_id = selected_chat_obj.get("department_id", 0)
                dept_name = selected_chat_obj.get("department_name", "Админ")
            else:
                client_id = 0
                dept_name = "Админ"
            
            st.info(f"👥 Подразделение: {dept_name}")
        else:
            chat_id = ""
            client_id = 0
            st.warning("Добавьте чат выше.")
        
        if st.button("Создать вакансию", key="btn_create_vacancy"):
            if new_title.strip() and chat_id:
                success, _ = create_vacancy(new_title.strip(), chat_id, client_id)
                if success:
                    st.success(f"Создано: {new_title}")
                    st.rerun()
                else:
                    st.error("Название уже существует.")
            else:
                st.warning("Заполните поля.")

    vacancies = load_vacancies()
    if not vacancies:
        st.info("Нет вакансий. Создайте первую.")
    else:
        st.divider()
        st.subheader("📋 Список вакансий")
        for i, vac in enumerate(vacancies):
            col_a, col_b, col_c = st.columns([4,2,1])
            col_a.write(f"**{vac['title']}**")
            col_b.caption(f"Chat ID: {vac.get('chat_id', '—')}")
            if col_c.button("🗑️", key=f"del_vac_{i}"):
                delete_vacancy(vac['title'])
                st.rerun()
        
        st.divider()
        selected_title = st.selectbox("Выберите вакансию для работы", [v["title"] for v in vacancies])
        selected_vacancy = next(v for v in vacancies if v["title"] == selected_title)

        with st.expander("📄 Документы вакансии (редактирование)"):
            docs = selected_vacancy.get("documents", {})
            new_profile = st.text_area("Профиль", value=docs.get("profile", ""), height=200)
            new_vac_text = st.text_area("Текст вакансии", value=docs.get("vacancy_text", ""), height=200)
            new_questions = st.text_area("Вопросы", value=docs.get("questions", ""), height=200)
            new_keywords = st.text_area("Ключевые слова", value=docs.get("keywords", ""), height=100)
            if st.button("💾 Сохранить документы"):
                update_vacancy_docs(selected_title, {
                    "profile": new_profile,
                    "vacancy_text": new_vac_text,
                    "questions": new_questions,
                    "keywords": new_keywords
                })
                st.success("Сохранено!")
                st.rerun()

        st.markdown(f"### 🔍 Кандидаты для «{selected_title}»")
        if selected_vacancy["candidates"]:
            for idx, cand in enumerate(selected_vacancy["candidates"]):
                with st.expander(f"👤 {cand.get('name', 'Без имени')}"):
                    col1, col2 = st.columns([4,1])
                    with col1:
                        st.markdown(f"**Резюме:** {cand.get('resume_link', '—')}")
                        st.markdown(f"**Видео:** {cand.get('video_link', '—')}")
                        transcript = st.text_area("Расшифровка собеседования", value=cand.get('transcript', ''), key=f"trans_{idx}", height=100)
                        if transcript != cand.get('transcript', ''):
                            selected_vacancy["candidates"][idx]['transcript'] = transcript
                            save_vacancies(vacancies)
                            st.success("Расшифровка сохранена!")
                        if cand.get('ai_score') is not None:
                            st.metric("Оценка ИИ", f"{cand['ai_score']}/10")
                            if cand.get('ai_comment'):
                                st.info(cand['ai_comment'])
                    with col2:
                        if st.button("🤖 Оценить", key=f"eval_{idx}"):
                            with st.spinner("Оценка..."):
                                resume_text = extract_text_from_pdf_url(cand.get('resume_link', '')) or ""
                                eval_result = evaluate_candidate_with_ai(resume_text, cand.get('transcript', ''), selected_title)
                                selected_vacancy["candidates"][idx]["ai_score"] = eval_result.get("score", 0)
                                selected_vacancy["candidates"][idx]["ai_comment"] = eval_result.get("comment", "")
                                selected_vacancy["candidates"][idx]["ai_strengths"] = eval_result.get("strengths", [])
                                selected_vacancy["candidates"][idx]["ai_weaknesses"] = eval_result.get("weaknesses", [])
                                save_vacancies(vacancies)
                                st.success("Оценка завершена!")
                                st.rerun()
                        hr_comment = st.text_area("Комментарий рекрутера", value=cand.get('hr_comment', ''), key=f"hr_{idx}", height=50)
                        if hr_comment != cand.get('hr_comment', ''):
                            selected_vacancy["candidates"][idx]['hr_comment'] = hr_comment
                            save_vacancies(vacancies)
                            st.success("Комментарий сохранён!")
                        if st.button("🗑️ Удалить", key=f"del_cand_{idx}"):
                            selected_vacancy["candidates"].pop(idx)
                            save_vacancies(vacancies)
                            st.rerun()
        else:
            st.info("Нет кандидатов.")

        st.markdown("### ➕ Добавить кандидата")
        with st.form("add_candidate", clear_on_submit=True):
            vacancies_list = load_vacancies()
            vacancy_titles = [v["title"] for v in vacancies_list]
            selected_vacancy_title = st.selectbox("Вакансия", vacancy_titles, key="cand_vacancy")
            selected_vacancy_obj = next(v for v in vacancies_list if v["title"] == selected_vacancy_title)
            name = st.text_input("ФИО")
            resume_link = st.text_input("Ссылка на резюме")
            video_link = st.text_input("Ссылка на видео")
            hr_comment = st.text_area("Комментарий")
            submitted = st.form_submit_button("Добавить и уведомить")
            if submitted:
                if name.strip():
                    new_cand = {
                        "vacancy_id": selected_vacancy_obj["id"],
                        "name": name.strip(),
                        "resume_link": resume_link.strip(),
                        "video_link": video_link.strip(),
                        "transcript": "",
                        "hr_comment": hr_comment.strip(),
                        "client_status": "wait",
                        "client_comment": "",
                        "office_interview_date": "",
                        "ai_score": None,
                        "ai_comment": "",
                        "ai_strengths": [],
                        "ai_weaknesses": []
                    }
                    selected_vacancy["candidates"].append(new_cand)
                    save_vacancies(vacancies)
                    bot_token = os.getenv("TELEGRAM_BOT_TOKEN")
                    chat_id = selected_vacancy.get("chat_id")
                    if bot_token and chat_id:
                        msg = f"🆕 <b>Новый кандидат</b>\n🏢 {selected_title}\n👤 {name.strip()}"
                        if resume_link.strip():
                            msg += f"\n📄 <a href='{resume_link.strip()}'>Резюме</a>"
                        if video_link.strip():
                            msg += f"\n🎥 <a href='{video_link.strip()}'>Видео</a>"
                        send_telegram_message(bot_token, chat_id, msg)
                    st.success("Кандидат добавлен!")
                    st.rerun()
                else:
                    st.error("Введите ФИО")

# ---------- ВКЛАДКА 5: ИНСТРУКЦИИ ----------
with tab5:
    st.header("📖 Инструкции по работе с HR-помогатором")
    st.markdown("""
    Добро пожаловать! Это приложение помогает HR-специалистам создавать документы для вакансий на основе расшифровки разговора с заказчиком, управлять воронкой кандидатов и отправлять уведомления в Telegram.
    
    **Основные возможности:**
    - Загрузка аудио/видео или текстового файла с описанием вакансии → автоматическое создание профиля должности, текста вакансии, опросника с примерами ответов и ключевых слов для поиска.
    - Редактирование сгенерированных документов и доработка через ИИ.
    - Сохранение документов в вакансии, создание новых вакансий с привязкой к Telegram-группе.
    - Воронка кандидатов: добавление кандидатов, оценка ИИ, уведомления заказчика в Telegram.
    - Экспорт готового пакета документов в Word и PDF.
    - История генераций: все ранее созданные пакеты сохраняются и доступны для повторной загрузки.
    """)
    with st.expander("🎤 1. Расшифровка файлов и генерация документов"):
        st.markdown("""
        **1. Загрузите файл**
        - Перейдите на вкладку **«Расшифровка»**.
        - Выберите источник: **«Аудио/видео»** или **«Готовый файл»**.
        - Нажмите **«Расшифровать»** или просто загрузите файл.

        **2. Отредактируйте текст (при необходимости)**
        - После загрузки появится текстовое поле. Вы можете исправить ошибки.
        - Нажмите **«Зафиксировать правки»**.

        **3. Сгенерируйте документы**
        - Перейдите на вкладку **«Генерация»**.
        - Нажмите **«✨ Сгенерировать всё»**.
        - Результат появится во вкладке **«Результаты»**.

        **4. Доработка через ИИ**
        - Введите дополнительные указания и нажмите **«🔄 Доработать»**.
        """)
    with st.expander("💾 2. Сохранение и обновление документов в вакансии"):
        st.markdown("""
        **Для сохранения в существующую вакансию:**
        - Во вкладке **«Генерация»** выберите вакансию и нажмите **«Обновить документы»**.

        **Для создания новой вакансии:**
        - Разверните **«➕ Создать новую вакансию»**, заполните поля и нажмите **«Создать и сохранить»**.
        """)
    with st.expander("🤖 3. Telegram: создание чатов и получение chat_id"):
        st.markdown("""
        **Как получить chat_id:**
        - После создания группы добавьте бота и отправьте любое сообщение.
        - Откройте: `https://api.telegram.org/bot<ВАШ_ТОКЕН>/getUpdates`
        - Найдите в ответе `"chat":{"id":-1001234567890,...}` и скопируйте число.
        """)
    with st.expander("👥 4. Работа с кандидатами в воронке"):
        st.markdown("""
        **Добавление кандидата:**
        - Выберите вакансию, заполните форму и нажмите **«Добавить и уведомить»**.

        **Оценка кандидата ИИ:**
        - После добавления кандидата нажмите **«🤖 Оценить»**.
        """)
    with st.expander("📎 5. Экспорт в Word и PDF"):
        st.markdown("""
        Во вкладке **«Результаты»** есть три кнопки для скачивания документов в форматах JSON, Word и PDF.
        """)
    with st.expander("📜 6. История генераций"):
        st.markdown("""
        Вкладка **«История»** сохраняет все когда‑либо созданные пакеты документов. Вы можете загрузить любой из них.
        """)
    st.info("💡 **Подсказка:** Для работы с аудио нужен установленный `ffmpeg`.")

# ---------- ВКЛАДКА 6: ИСТОРИЯ ГЕНЕРАЦИЙ ----------
with tab6:
    st.header("📜 История генераций")
    st.caption("Все ранее созданные пакеты документов. Вы можете загрузить любой из них для просмотра, доработки или экспорта.")
    
    index = get_history_index()
    if not index:
        st.info("История пуста. Сгенерируйте документы во вкладке «Генерация», чтобы они появились здесь.")
    else:
        for i, rec in enumerate(index):
            with st.expander(f"📄 {rec['datetime']} – {rec['title'] or 'Без названия'}"):
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.markdown(f"**Время:** {rec['datetime']}")
                    st.markdown(f"**Должность:** {rec['title']}")
                    if rec['vacancy_title']:
                        st.markdown(f"**Связанная вакансия:** {rec['vacancy_title']}")
                    st.markdown(f"**Превью текста вакансии:**")
                    st.text(rec['preview'] if rec['preview'] else "(нет текста)")
                with col2:
                    if st.button("📂 Загрузить", key=f"load_{i}"):
                        data = load_generation_from_history(rec['filename'])
                        if data:
                            st.session_state.generated = data
                            st.success(f"Загружен пакет от {rec['datetime']}! Перейдите во вкладку «Результаты» или «Генерация».")
                            st.rerun()
                        else:
                            st.error("Ошибка загрузки файла.")
                    data_for_export = load_generation_from_history(rec['filename'])
                    if data_for_export:
                        st.download_button(
                            "📥 Скачать JSON",
                            data=json.dumps(data_for_export, ensure_ascii=False, indent=2),
                            file_name=rec['filename'],
                            mime="application/json",
                            key=f"export_hist_{i}"
                        )
                    if st.button("🗑️ Удалить", key=f"del_{i}"):
                        if delete_generation_from_history(rec['filename']):
                            st.success("Пакет удалён из истории.")
                            st.rerun()
                        else:
                            st.error("Ошибка удаления.")