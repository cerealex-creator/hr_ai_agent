# hri_full_v1.py — HR-ассистент с динамическими подразделениями
# Запуск: streamlit run hri_full_v1.py

import streamlit as st
import os
import yaml
import json
import time
import requests
import boto3
import subprocess
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook
import PyPDF2
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import mm
import tempfile
from datetime import datetime
import glob
import re
import logging
from urllib.parse import quote
from corporate_ui import apply_corporate_ui
from eval_ui import render_ai_score_badge
from vacancy_tab import render_vacancy_tab
from stats_tab import render_stats_tab
from models import migrate_candidate
from vacancy_display import format_vacancy_search_period
from warranty import default_warranty, migrate_vacancy_warranty

# ============================================================
# ОСНОВНЫЕ КОНСТАНТЫ
# ============================================================
CLIENTS = {0: "Админ"}  # Будет дополняться из departments.json

# ============================================================
# УПРАВЛЕНИЕ ПОДРАЗДЕЛЕНИЯМИ
# ============================================================
DEPARTMENTS_FILE = "data/departments.json"

def load_departments():
    """Загружает список подразделений из файла"""
    if not os.path.exists(DEPARTMENTS_FILE):
        # Создаём с дефолтными подразделениями
        default = {
            "departments": [
                {"id": 1, "name": "Маркетинг", "slug": "marketing"},
                {"id": 2, "name": "Продажи", "slug": "sales"},
                {"id": 3, "name": "Бухгалтерия", "slug": "accounting"},
                {"id": 4, "name": "Склад", "slug": "warehouse"},
                {"id": 5, "name": "Логистика", "slug": "logistics"},
                {"id": 6, "name": "Руководители", "slug": "managers"},
                {"id": 99, "name": "Тестировочный", "slug": "test"}
            ]
        }
        save_departments(default)
        return default["departments"]
    with open(DEPARTMENTS_FILE, "r", encoding="utf-8") as f:
        data = json.load(f)
        return data.get("departments", [])

def save_departments(departments_data):
    """Сохраняет список подразделений в файл"""
    os.makedirs("data", exist_ok=True)
    with open(DEPARTMENTS_FILE, "w", encoding="utf-8") as f:
        json.dump(departments_data, f, ensure_ascii=False, indent=2)

def add_department(name):
    """Добавляет новое подразделение и возвращает его ID"""
    departments = load_departments()
    new_id = max([d["id"] for d in departments], default=0) + 1
    slug = name.lower().replace(" ", "_")
    departments.append({"id": new_id, "name": name, "slug": slug})
    save_departments({"departments": departments})
    return new_id

def get_department_by_id(dept_id):
    """Возвращает подразделение по ID"""
    departments = load_departments()
    for d in departments:
        if d["id"] == dept_id:
            return d
    return None

# ============================================================
# ЗАГРУЗКА КОНФИГУРАЦИИ И КЛЮЧЕЙ
# ============================================================
load_dotenv(os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env"), override=True)

def load_config(config_path="hri_full_v1_config.yaml"):
    with open(config_path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)

config = load_config()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(name)s: %(message)s",
)

api_key_env = config["model"].get("api_key_env", "PROXYAPI_KEY")
api_key = os.getenv(api_key_env)
base_url = config["model"].get("base_url", "https://api.proxyapi.ru/openai/v1")

client = OpenAI(
    base_url=base_url,
    api_key=api_key
)

# ============================================================
# ВСТРОЕННЫЕ ОБРАЗЦЫ ДОКУМЕНТОВ
# ============================================================
EXAMPLE_PROFILE = """
Профиль должности (указать должность)
Описание
Комментарии
Общие данные
Подразделение

Непосредственный руководитель сотрудника

ЦКП (ценный конечный продукт - за что именно платятся деньги) должности

Основные обязанности

Анкетные данные
Предыдущие места работы для нужного опыта

Предыдущие должности для нужного опыта

Образование: 

Пол:  

Возраст:  

Семейное положение, дети:

Место жительства:

Прочие требования: 

Стоп факторы: 
Например, очное обучение; неготовность к офисному формату

Hard Skills, без которых кандидат не сможет работать:
Навык, умение или знание: 
Как проявляется на стадии собеседований (что хотим услышать):

Hard Skills, которые толковый кандидат может освоить в первые 1-6 недель работы:

Личные качества
Качество/Как проявляется
Как увидим на стадии собеседований?

Условия работы:
Формат, режим
Оклад (несгораемая цифра)
Переменная часть - бонус
Длительность испытательного срока
Зарплатные условия на период испытательного срока
"""

EXAMPLE_FILLED_PROFILE = """
Профиль должности: Младший менеджер по маркетплейсам (Wildberries)

Подразделение: Отдел маркетинга

Непосредственный руководитель: Руководитель отдела маркетинга

ЦКП должности: Выполненные в срок и качественно задачи по наполнению контентом карточек товаров

Основные обязанности:
- Создание карточек товаров (работа в паре со старшим менеджером)
- Наполнение карточки описанием и характеристиками
- Работа с баллами за отзывы (своевременный запуск)
- Работа с рекомендациями продавца в карточках товаров
- Внесение изменений в карточки товаров по запросу от других сотрудников
- Ответы на отзывы и вопросы (в рамках помощи ответственному специалисту)
- SEO в карточках товаров (сбор ключевых запросов, выделение ключей)

Анкетные данные:
- Образование: желательно высшее/среднее специальное
- Пол: лучше женский
- Возраст: от 25 лет, желательно 30-45 лет
- Семейное положение: если женщина, то без детей, либо дети школьного возраста
- Место жительства: СПб Янино, не более 1 часа дороги до офиса

Стоп-факторы:
- Частая смена работы (несколько последних мест работы не более 3 месяцев)
- Неготовность к офисному формату работы

Hard Skills (обязательные):
- Знание специфики работы с маркетплейсами (может ответить на проверочные вопросы)
- Опыт любой работы с карточками товаров (выполняет простую тестовую задачу)

Hard Skills (желательные, освоит за 1-6 недель):
- Опыт работы с контентом в карточках товаров
- Знание возможностей массового редактирования в ЛК ВБ
- Знание 1С на базовом уровне
- Опыт работы с SEO в текстах карточек

Личные качества:
- Высокая скорость и лояльность к рутинным задачам (делает задачи в срок, быстро соображает)
- Средний уровень предпринимателя (творчество как вспомогательный инструмент)
- Уровень Интегратора ниже среднего (не стремится понравиться всем)
- Администратор уровне выше среднего (уделяет внимание важным деталям)

Условия работы:
- Формат: Офисный (офис в СПб на м. Ладожская)
- Режим: 09:00-18:00
- Оклад: 80 000 руб. после испытательного срока
- Испытательный срок: 1-3 месяца
- Зарплата на испытательный срок: 70 000 руб.
"""

EXAMPLE_VACANCY = """
Младший менеджер по маркетплейсам (Wildberries)
от 75 000 до 85 000 ₽ за месяц, до вычета налогов
Торговый дом
Ладожская, Санкт-Петербург, проспект Энергетиков, 4к1

Выплаты: два раза в месяц
Опыт работы: не требуется
Оформление: Трудовой договор
График: 5/2
Формат работы: на месте работодателя

YourBox — крупный бренд обуви из Санкт-Петербурга с 10-летней историей.

Основные задачи:
— Создание и редактирование карточек товаров на платформе;
— Проведение простых А/Б тестов для повышения конверсии;
— Работа с отзывами и баллами за отзывы;
— Создание баннерных рекламных кампаний и управление ими (WB.Media);
— Организация фотосессий товаров;
— Работа с фото- и видеоконтентом;
— SEO-оптимизация карточек товаров.

Требования к кандидату:
— Опыт работы с карточками товаров на Wildberries;
— Внимательность к деталям;
— Ответственный подход к задачам.

Мы предлагаем:
— современное оборудованное пространство;
— обучение и поддержку;
— возможности для роста;
— возможность приобрести обувь по себестоимости.
"""

EXAMPLE_QUESTIONNAIRE = """
№   Вопрос                                                                                         Желательный результат
1    Почему вы сейчас ищете новую работу?                                                         Объективная причина или стремление к росту.
2    Что для вас важно в работе, и какие условия вы ищете?                                        Интерес к задачам компании, без излишнего фокуса на комфорт.
3    Какой у вас опыт работы с карточками товаров на Wildberries?                                 Опыт от 1 года, работа с 30+ артикулами.
4    Опишите ситуацию, когда вы заметили ошибку в карточке товара и как её исправили.            Конкретный пример, внимательность.
5    Как вы оцениваете эффективность рекламной кампании на Wildberries?                          Знание метрик CTR, CR, ACoS.
6    Как вы организуете свою работу при ведении большого количества товаров?                     Приоритизация, планирование.
7    Где вы учились работе с маркетплейсами?                                                      Курсы, самообучение, практика.
8    Итог:                                                                                       Проходной балл.
"""

EXAMPLE_FILLED_QUESTIONNAIRE = """
1. Основной: «Расскажите, почему сейчас рассматриваете смену работы? Что для вас важно в новом месте?»
   Уточняющие: «Что именно не устраивает на текущем месте?», «Какие условия для вас принципиальны?»
   Проверяет: мотивация, ожидания по условиям | Категория: experience

2. Основной: «Был ли у Вас опыт работы с карточками товаров на маркетплейсах? Как именно Вы это делали — расскажите на конкретном примере.»
   Уточняющие: «Сколько артикулов вели одновременно?», «Какие задачи выполняли сами, а что делегировали?»
   Проверяет: обязательный hard skill «опыт с карточками товаров» | Категория: hard_skills
"""

QUESTIONNAIRE_GENERATION_RULES = """
**Требования к опроснику для первичного собеседования:**

**Количество и плотность:**
- Оптимально 6–8 ОСНОВНЫХ вопросов. Максимум 10 — только если критичные невыясненные моменты профиля/резюме невозможно закрыть в 8 вопросах.
- Не дублируй смысл: если два вопроса проверяют одно и то же — объедини в один основной + 1–3 уточняющих.
- Не выноси мелочи и второстепенные детали в отдельные основные вопросы — перенеси в уточняющие.
- Каждый основной вопрос привязан к конкретному требованию профиля (поле проверяет_требование).

**ОБЯЗАТЕЛЬНЫЕ основные вопросы (включи все четыре, формулировки адаптируй под должность):**
1. Причина поиска и ухода с последнего места (или переход с фриланса/самозанятости): почему сейчас ищет работу, что произошло. Уточнения: длительный пробел в резюме — чем занимался; при уходе с фриланса — почему решил работать в штате/компании.
2. Что вдохновляет на работе и повышает желание работать (мотивация, ценные условия, что «заряжает»).
3. Что расстраивает на работе и заставляет задуматься об уходе (триггеры недовольства, «красные линии»).
4. Обратная связь от прошлых работодателей: готовность дать контакты для рекомендаций ИЛИ «хорошо ли расстались, что могли бы сказать о Вас предыдущие руководители?».

**Остальные вопросы (в пределах 6–8 суммарно с обязательными):**
- Обязательные hard skills, ключевой релевантный опыт, 1–2 soft skills из профиля.
- Стиль беседы, не допрос: «Был ли у Вас опыт ...? Расскажите на примере», «Расскажите о ситуации, когда ...».
- К каждому основному вопросу — 1–3 уточняющих (уточняющие_вопросы), если ответ общий или неполный.
- категория: hard_skills / soft_skills / experience / motivation / reliability.
- пример_ответа — реалистичный ответ сильного кандидата с конкретикой.

Опросник должен позволить по расшифровке интервью оценить профессиональное соответствие И лояльность, адекватность, управляемость кандидата.
"""

# ============================================================
# ФУНКЦИИ ДЛЯ РАБОТЫ С ФАЙЛАМИ И ТЕКСТОМ
# ============================================================
def extract_text(uploaded_file):
    filename = uploaded_file.name.lower()
    if filename.endswith(".txt"):
        return uploaded_file.read().decode("utf-8")
    elif filename.endswith(".docx"):
        doc = Document(uploaded_file)
        return "\n".join([p.text for p in doc.paragraphs])
    elif filename.endswith(".xlsx"):
        wb = load_workbook(BytesIO(uploaded_file.getvalue()), read_only=True, data_only=True)
        sheet = wb.active
        rows = []
        for row in sheet.iter_rows(values_only=True):
            if not any(cell is not None and str(cell).strip() for cell in row):
                continue
            row_text = " | ".join(str(cell).strip() if cell is not None else "" for cell in row)
            rows.append(row_text)
        wb.close()
        return "\n".join(rows)
    elif filename.endswith(".pdf"):
        reader = PyPDF2.PdfReader(uploaded_file)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    else:
        raise ValueError("Неподдерживаемый формат файла")

# Улучшенная функция получения прямой ссылки с Яндекс.Диска 
def get_direct_yandex_link(public_url):
    """Преобразует публичную ссылку Яндекс.Диска в прямую ссылку на скачивание."""
    if not public_url:
        return None
    if "disk.yandex.ru" in public_url or "yadi.sk" in public_url:
        from resume_ai import get_yandex_download_url

        return get_yandex_download_url(public_url)
    return public_url

# Улучшенная функция извлечения текста из PDF 
def extract_text_from_pdf_url(pdf_url):
    """
    Скачивает PDF по ссылке с Яндекс.Диска и извлекает текст.
    """
    if not pdf_url:
        print("⚠️ PDF URL пустой")
        return ""
        
    try:
        print(f"📥 Получаем прямую ссылку для PDF...")
        download_url = get_direct_yandex_link(pdf_url)
        if not download_url:
            print("❌ Не удалось получить прямую ссылку")
            return ""
        
        print(f"📥 Скачиваем файл...")
        response = requests.get(download_url, timeout=60)
        if response.status_code == 200:
            print(f"✅ Файл скачан, размер: {len(response.content)} байт")
            if not response.content.lstrip().startswith(b"%PDF"):
                content_type = response.headers.get("Content-Type", "")
                print(f"⚠️ Это не PDF (Content-Type: {content_type})")
                return ""
            pdf_file = BytesIO(response.content)
            try:
                reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                        print(f"  Страница {i+1}: {len(page_text)} символов")
                result = text.strip()
                print(f"✅ Извлечено текста: {len(result)} символов")
                return result
            except Exception as e:
                print(f"❌ Ошибка чтения PDF: {e}")
                return ""
        else:
            print(f"❌ Ошибка скачивания PDF: статус {response.status_code}")
            return ""
    except Exception as e:
        print(f"❌ Ошибка извлечения текста из PDF: {e}")
        import traceback
        traceback.print_exc()
        return ""

# Новая функция транскрипции видео по ссылке 
def transcribe_video_from_link(video_link):
    """
    Скачивает видео/аудио по ссылке, конвертирует и возвращает текст.
    Исправленная версия с надёжной конвертацией.
    """
    if not video_link:
        print("⚠️ Ссылка на видео пустая")
        return ""
    
    tmp_path = None
    pcm_path = None
    
    try:
        print(f"🎥 Начинаем расшифровку видео...")
        print(f"🔗 Ссылка: {video_link[:50]}...")
        
        # Получаем прямую ссылку
        print(f"📥 Получаем прямую ссылку...")
        direct_url = get_direct_yandex_link(video_link)
        if not direct_url:
            print("❌ Не удалось получить прямую ссылку на видео")
            return ""
        
        print(f"✅ Прямая ссылка получена")
        
        # Скачиваем файл
        print(f"📥 Скачиваем файл (это может занять время)...")
        response = requests.get(direct_url, timeout=300, stream=True)
        if response.status_code != 200:
            print(f"❌ Ошибка скачивания видео: статус {response.status_code}")
            return ""
        
        # Определяем расширение
        content_type = response.headers.get('content-type', '')
        print(f"📄 Content-Type: {content_type}")
        
        ext = '.mp4'
        if 'webm' in content_type or '.webm' in video_link:
            ext = '.webm'
        elif 'mp3' in content_type or '.mp3' in video_link:
            ext = '.mp3'
        elif 'wav' in content_type or '.wav' in video_link:
            ext = '.wav'
        elif 'ogg' in content_type or '.ogg' in video_link:
            ext = '.ogg'
        
        # Сохраняем во временный файл
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp.write(response.content)
            tmp_path = tmp.name
        
        file_size = os.path.getsize(tmp_path)
        print(f"✅ Файл скачан: {tmp_path}, размер: {file_size / (1024*1024):.1f} МБ")

        # Одна конвертация в PCM внутри transcribe_speechkit_cloud
        print("🎙️ Расшифровка через Яндекс SpeechKit...")
        text = transcribe_speechkit_cloud(tmp_path)
        print(f"✅ Расшифровка завершена! Текст: {len(text)} символов")

        return text.strip()
        
    except subprocess.CalledProcessError as e:
        print(f"❌ Ошибка ffmpeg: {e}")
        print(f"stderr: {e.stderr.decode() if e.stderr else 'нет данных'}")
        return ""
    except Exception as e:
        print(f"❌ Ошибка транскрипции: {e}")
        import traceback
        traceback.print_exc()
        return ""
    finally:
        # Очищаем временные файлы
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)
            print(f"🗑️ Временный файл удален")
        if pcm_path and os.path.exists(pcm_path):
            os.unlink(pcm_path)
            print(f"🗑️ PCM файл удален")

# ============================================================
# ФУНКЦИИ ДЛЯ ЯНДЕКС ОБЛАКА
# ============================================================
def _validate_speechkit_config(*, bucket=None, access_key=None, secret_key=None, api_key=None):
    missing = []
    if not (bucket or "").strip():
        missing.append("YANDEX_BUCKET_NAME")
    if not (access_key or "").strip():
        missing.append("YANDEX_ACCESS_KEY_ID")
    if not (secret_key or "").strip():
        missing.append("YANDEX_SECRET_ACCESS_KEY")
    if api_key is not None and not (api_key or "").strip():
        missing.append("YANDEX_API_KEY")
    if missing:
        raise RuntimeError(
            "Не настроен Яндекс SpeechKit / Object Storage. Не хватает: "
            + ", ".join(missing)
            + ". Проверьте файл `.env` на сервере."
        )


def upload_to_s3_and_get_url(local_path, bucket, access_key, secret_key):
    """
    Загружает локальный файл в Yandex Object Storage и возвращает публичную ссылку.
    БЕЗ callback — чтобы избежать ошибки NoSessionContext в Streamlit.
    """
    import boto3
    from boto3 import client
    
    _validate_speechkit_config(
        bucket=bucket, access_key=access_key, secret_key=secret_key, api_key=None
    )

    try:
        # Создаем клиент НАПРЯМУЮ, без session.Session()
        s3_client = client(
            's3',
            endpoint_url='https://storage.yandexcloud.net',
            aws_access_key_id=access_key,
            aws_secret_access_key=secret_key
        )
        
        object_name = os.path.basename(local_path)
        if not os.path.exists(local_path):
            raise RuntimeError(f"Файл для загрузки не найден: {local_path}")
        
        # Загружаем файл БЕЗ callback!
        s3_client.upload_file(local_path, bucket, object_name)
        
        # Формируем прямую ссылку
        url = f"https://storage.yandexcloud.net/{bucket}/{object_name}"
        return url
        
    except Exception as e:
        raise RuntimeError(
            "Ошибка загрузки файла в Yandex Object Storage. "
            "Проверьте, что бакет существует и ключи имеют права `storage.editor`. "
            f"Детали: {type(e).__name__}: {e}"
        ) from e

def convert_to_pcm(input_path, output_path=None):
    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + "_speechkit.pcm"
    try:
        subprocess.run(
            [
                "ffmpeg",
                "-y",
                "-i",
                input_path,
                "-acodec",
                "pcm_s16le",
                "-ac",
                "1",
                "-ar",
                "16000",
                "-f",
                "s16le",
                output_path,
            ],
            check=True,
            capture_output=True,
        )
    except FileNotFoundError as e:
        raise RuntimeError(
            "Не найден `ffmpeg`. Установите ffmpeg и повторите попытку."
        ) from e
    except subprocess.CalledProcessError as e:
        stderr = (e.stderr or b"").decode("utf-8", errors="ignore").strip()
        raise RuntimeError(
            "Ошибка конвертации через ffmpeg. "
            + (f"stderr: {stderr}" if stderr else "Проверьте формат файла.")
        ) from e
    return output_path

def recognize_long_audio(audio_url, api_key):
    _validate_speechkit_config(
        bucket="ok", access_key="ok", secret_key="ok", api_key=api_key
    )
    response = requests.post(
        "https://transcribe.api.cloud.yandex.net/speech/stt/v2/longRunningRecognize",
        headers={"Authorization": f"Api-Key {api_key}"},
        json={
            "config": {
                "specification": {
                    "languageCode": "ru-RU",
                    "profanityFilter": "false",
                    "audioEncoding": "LINEAR16_PCM",
                    "sampleRateHertz": 16000,
                    "audioChannelCount": 1
                }
            },
            "audio": {"uri": audio_url}
        }
    )
    if response.status_code != 200:
        raise RuntimeError(
            f"SpeechKit: ошибка запроса ({response.status_code}). {response.text}"
        )
    operation = response.json()
    operation_id = operation["id"]
    while True:
        time.sleep(5)
        status_response = requests.get(
            f"https://operation.api.cloud.yandex.net/operations/{operation_id}",
            headers={"Authorization": f"Api-Key {api_key}"}
        )
        if status_response.status_code != 200:
            raise RuntimeError(
                f"SpeechKit: ошибка статуса операции ({status_response.status_code}). {status_response.text}"
            )
        status_data = status_response.json()
        if status_data.get("done"):
            if "error" in status_data:
                raise RuntimeError(f"SpeechKit: ошибка распознавания: {status_data['error']}")
            full_text = ""
            for chunk in status_data["response"]["chunks"]:
                full_text += chunk["alternatives"][0]["text"] + " "
            return full_text.strip()

# ============================================================
# УПРАВЛЕНИЕ ДАННЫМИ: ВАКАНСИИ, ЧАТЫ, ИСТОРИЯ
# ============================================================
VACANCIES_FILE = "data/vacancies_db.json"
CHATS_FILE = "data/chats_db.json"
HISTORY_DIR = "data/history"
HISTORY_INDEX_FILE = os.path.join(HISTORY_DIR, "index.json")

def ensure_history_dir():
    os.makedirs(HISTORY_DIR, exist_ok=True)


def sanitize_history_slug(text):
    """Безопасное имя файла: убирает / и другие символы, ломающие путь."""
    if not text or not str(text).strip():
        return "unknown"
    slug = str(text).strip().replace(" ", "_")
    slug = re.sub(r'[/\\:*?"<>|]', "_", slug)
    slug = re.sub(r"_+", "_", slug).strip("_.")
    return slug[:120] if slug else "unknown"


def save_generation_to_history(generated_data, transcript_text=None, vacancy_title=None):
    ensure_history_dir()
    timestamp = int(time.time())
    dt_str = time.strftime("%Y%m%d_%H%M%S", time.localtime(timestamp))
    title = sanitize_history_slug(generated_data.get("должность", "unknown"))
    if vacancy_title:
        title = sanitize_history_slug(vacancy_title)
    filename = f"{dt_str}_{title}.json"
    filepath = os.path.join(HISTORY_DIR, filename)
    
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(generated_data, f, ensure_ascii=False, indent=2)
    
    if os.path.exists(HISTORY_INDEX_FILE):
        with open(HISTORY_INDEX_FILE, "r", encoding="utf-8") as f:
            index = json.load(f)
    else:
        index = []
    
    record = {
        "timestamp": timestamp,
        "datetime": dt_str,
        "filename": filename,
        "title": generated_data.get("должность", ""),
        "preview": generated_data.get("текст_вакансии", "")[:100] + "..." if generated_data.get("текст_вакансии") else "",
        "vacancy_title": vacancy_title or "",
        "has_transcript": bool(transcript_text)
    }
    index.append(record)
    index.sort(key=lambda x: x["timestamp"], reverse=True)
    with open(HISTORY_INDEX_FILE, "w", encoding="utf-8") as f:
        json.dump(index, f, ensure_ascii=False, indent=2)
    return record

def load_generation_from_history(filename):
    filepath = os.path.join(HISTORY_DIR, filename)
    if os.path.exists(filepath):
        with open(filepath, "r", encoding="utf-8") as f:
            return json.load(f)
    return None

def delete_generation_from_history(filename):
    filepath = os.path.join(HISTORY_DIR, filename)
    if os.path.exists(filepath):
        os.remove(filepath)
    if os.path.exists(HISTORY_INDEX_FILE):
        with open(HISTORY_INDEX_FILE, "r", encoding="utf-8") as f:
            index = json.load(f)
        new_index = [rec for rec in index if rec["filename"] != filename]
        with open(HISTORY_INDEX_FILE, "w", encoding="utf-8") as f:
            json.dump(new_index, f, ensure_ascii=False, indent=2)
        return True
    return False

def get_history_index():
    if os.path.exists(HISTORY_INDEX_FILE):
        with open(HISTORY_INDEX_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

def load_chats():
    if not os.path.exists(CHATS_FILE):
        return []
    with open(CHATS_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

def save_chats(chats_list):
    os.makedirs("data", exist_ok=True)
    with open(CHATS_FILE, "w", encoding="utf-8") as f:
        json.dump(chats_list, f, ensure_ascii=False, indent=2)

def load_vacancies():
    if not os.path.exists(VACANCIES_FILE):
        return []
    try:
        from vacancy_store import load_vacancies as _load_vacancies_store
        vacancies = _load_vacancies_store().get("vacancies", [])
    except json.JSONDecodeError as e:
        st.error(f"Ошибка чтения файла vacancies_db.json: {e}. Данные повреждены.")
        # Создаем бэкап повреждённого файла
        backup_name = f"{VACANCIES_FILE}.broken_{int(time.time())}.json"
        try:
            os.rename(VACANCIES_FILE, backup_name)
            st.warning(f"Повреждённый файл сохранён как {backup_name}. Будет создан новый файл.")
        except:
            pass
        backup_files = glob.glob(f"{VACANCIES_FILE}.backup_*")
        if backup_files:
            latest_backup = max(backup_files, key=os.path.getctime)
            try:
                with open(latest_backup, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    vacancies = data.get("vacancies", [])
                st.info(f"Загружен бэкап от {latest_backup}")
                return vacancies
            except:
                pass
        vacancies = []
        save_vacancies(vacancies)
        return vacancies

    migrated = False
    for v in vacancies:
        if "documents" not in v:
            v["documents"] = {
                "profile": "",
                "vacancy_text": "",
                "questions": "",
                "keywords": "",
                "notes": ""
            }
            migrated = True
        if "active" not in v:
            v["active"] = True
            migrated = True
        if "created_at" not in v:
            v["created_at"] = datetime.now().isoformat()
            migrated = True
        if "closed_at" not in v:
            v["closed_at"] = None
            migrated = True
        if "vacancy_summary" not in v:
            v["vacancy_summary"] = ""
            migrated = True
        if migrate_vacancy_warranty(v):
            migrated = True
        from vacancy_close import migrate_vacancy_close

        if migrate_vacancy_close(v):
            migrated = True
        from yandex_disk_ingest import migrate_vacancy_yandex_disk

        if migrate_vacancy_yandex_disk(v):
            migrated = True
        from vacancy_stats_filter import migrate_vacancy_is_test

        if migrate_vacancy_is_test(v):
            migrated = True
        for candidate in v.get("candidates", []):
            if "task_link" not in candidate:
                candidate["task_link"] = ""
                migrated = True
            if "office_interview_time" not in candidate:
                candidate["office_interview_time"] = ""
                migrated = True
            if "client_final_verdict" not in candidate:
                candidate["client_final_verdict"] = ""
                migrated = True
            if "ignore_flags" not in candidate or not isinstance(candidate.get("ignore_flags"), dict):
                candidate["ignore_flags"] = default_ignore_flags()
                migrated = True
            else:
                for flag_key, flag_val in default_ignore_flags().items():
                    if flag_key not in candidate["ignore_flags"]:
                        candidate["ignore_flags"][flag_key] = flag_val
                        migrated = True
            if "profile_checked" not in candidate:
                candidate["profile_checked"] = False
                migrated = True
            if "transcript" not in candidate:
                candidate["transcript"] = ""
                migrated = True
            if migrate_candidate(candidate, default_ignore_flags):
                migrated = True
    if migrated:
        save_vacancies(vacancies)
    return vacancies

def save_vacancies(vacancies_list):
    from vacancy_store import save_vacancies_list
    save_vacancies_list(vacancies_list)

def _empty_vacancy_documents():
    return {
        "profile": "",
        "vacancy_text": "",
        "questions": "",
        "keywords": "",
        "notes": "",
    }


def _next_vacancy_id(vacancies):
    if not vacancies:
        return 1
    return max(v.get("id", 0) for v in vacancies) + 1


def create_vacancy(title, chat_id, client_id=0, *, documents=None, show_portfolio_field=False, is_test=False):
    from telegram_notify import normalize_chat_id
    vacancies = load_vacancies()
    title = (title or "").strip()
    if not title:
        return False, "Введите название должности"
    if any(v.get("title") == title and v.get("active", True) for v in vacancies):
        return False, (
            "Уже есть активная вакансия с таким названием. "
            "Переместите предыдущую итерацию в архив или укажите другое название."
        )
    docs = documents if documents is not None else _empty_vacancy_documents()
    new_vacancy = {
        "id": _next_vacancy_id(vacancies),
        "title": title,
        "chat_id": normalize_chat_id(chat_id),
        "client_id": client_id,
        "active": True,
        "created_at": datetime.now().isoformat(),
        "closed_at": None,
        "vacancy_summary": "",
        "documents": docs,
        "candidates": [],
        "show_portfolio_field": bool(show_portfolio_field),
        "search_mode": "normal",
        "warranty_source_vacancy_id": None,
        "warranty": default_warranty(),
        "close_reason": None,
        "is_test": bool(is_test),
    }
    vacancies.append(new_vacancy)
    save_vacancies(vacancies)
    return True, new_vacancy


def create_vacancy_from_template(template_id, title, chat_id=None, client_id=None, *, is_test=False):
    import copy

    from vacancy_template_store import get_template

    template = get_template(template_id)
    if not template:
        return False, "Шаблон не найден"

    resolved_chat = chat_id if chat_id is not None else template.get("chat_id")
    resolved_client = client_id if client_id is not None else template.get("client_id", 0)
    documents = copy.deepcopy(template.get("documents") or _empty_vacancy_documents())
    return create_vacancy(
        title or template.get("title", ""),
        resolved_chat,
        resolved_client,
        documents=documents,
        is_test=is_test,
    )

def update_vacancy_docs_by_id(vacancy_id, docs_dict, *, replace_documents=False):
    """Обновляет documents у вакансии по id (надёжнее, чем по названию)."""
    vacancies = load_vacancies()
    for v in vacancies:
        if v.get("id") != vacancy_id:
            continue
        if "documents" not in v:
            v["documents"] = _empty_vacancy_documents()
        if replace_documents:
            notes = (v["documents"].get("notes") or "")
            v["documents"] = {**_empty_vacancy_documents(), "notes": notes}
        v["documents"].update(docs_dict)
        save_vacancies(vacancies)
        return True
    return False


def update_vacancy_docs(vacancy_title, docs_dict, *, vacancy_id=None):
    if vacancy_id is not None:
        return update_vacancy_docs_by_id(vacancy_id, docs_dict)
    vacancies = load_vacancies()
    matches = [v for v in vacancies if v.get("title") == vacancy_title]
    if len(matches) > 1:
        active = [v for v in matches if v.get("active", True)]
        if len(active) == 1:
            return update_vacancy_docs_by_id(active[0]["id"], docs_dict)
    for v in vacancies:
        if v["title"] == vacancy_title:
            if "documents" not in v:
                v["documents"] = {
                    "profile": "", "vacancy_text": "", "questions": "", "keywords": "", "notes": ""
                }
            v["documents"].update(docs_dict)
            save_vacancies(vacancies)
            return True
    return False

def delete_vacancy(vacancy_title):
    from vacancy_store import delete_vacancy_by_title

    ok, _ = delete_vacancy_by_title(vacancy_title)
    return ok

# ============================================================
# ФУНКЦИИ TELEGRAM
# ============================================================
def send_telegram_message(bot_token, chat_id, text):
    from telegram_notify import send_telegram_html
    ok, msg, _ = send_telegram_html(chat_id, text, bot_token=bot_token)
    return ok, msg

# ============================================================
# ОЦЕНКА КАНДИДАТА ИИ
# ============================================================
def default_ignore_flags():
    return {
        "ignore_age": False,
        "ignore_experience": False,
        "ignore_work_format": False,
        "ignore_office_distance": False,
        "ignore_hard_skills": False,
        "ignored_hard_skills_text": "",
    }


def vacancy_has_profile(vacancy):
    profile = vacancy.get("documents", {}).get("profile", "")
    if not profile or not str(profile).strip():
        return False
    text = str(profile).strip()
    if text.startswith("{"):
        try:
            parsed = json.loads(text)
            return bool(parsed)
        except json.JSONDecodeError:
            return False
    return len(text) >= 30


def get_vacancy_profile_text(vacancy):
    profile = vacancy.get("documents", {}).get("profile", "")
    if not profile:
        return ""
    text = str(profile).strip()
    if text.startswith("{"):
        try:
            parsed = json.loads(text)
            raw = (parsed.get("raw") or "").strip()
            if raw:
                return raw
            return json.dumps(parsed, ensure_ascii=False, indent=2)
        except json.JSONDecodeError:
            return text
    return text


def parse_ai_json_response(content):
    content = str(content or "").strip()
    content = re.sub(
        r"<(?:think|thinking|redacted_thinking)>.*?</(?:think|thinking|redacted_thinking)>",
        "",
        content,
        flags=re.DOTALL | re.IGNORECASE,
    ).strip()
    if "```json" in content:
        content = content.split("```json")[1].split("```")[0]
    elif "```" in content:
        content = content.split("```")[1].split("```")[0]
    content = content.strip()
    if not content.startswith("{") and not content.startswith("["):
        match = re.search(r"\{.*\}", content, re.DOTALL)
        if match:
            content = match.group(0)
    return json.loads(content.strip())


def build_ignore_flags_prompt(ignore_flags):
    flags = ignore_flags or default_ignore_flags()
    active = []
    if flags.get("ignore_age"):
        active.append("ignore_age — требования к возрасту")
    if flags.get("ignore_experience"):
        active.append("ignore_experience — требования к опыту")
    if flags.get("ignore_work_format"):
        active.append("ignore_work_format — формат работы (офис/удалёнка)")
    if flags.get("ignore_office_distance"):
        active.append("ignore_office_distance — расстояние до офиса")
    if flags.get("ignore_hard_skills"):
        skills_text = flags.get("ignored_hard_skills_text", "").strip()
        detail = f" ({skills_text})" if skills_text else ""
        active.append(f"ignore_hard_skills — конкретные hard skills{detail}")
    if not active:
        return "Активные флаги игнорирования: нет"
    return "Активные флаги игнорирования:\n- " + "\n- ".join(active)


SCORE_CATEGORY_RANGES = {
    0: (0, 29),
    1: (10, 39),
    2: (30, 79),
    3: (60, 89),
    4: (90, 100),
}

EVAL_CATEGORY_KEYS = ("hard_skills", "soft_skills", "experience")

EVAL_SYSTEM_PROMPT = """Ты — опытный HR-директор. Оцениваешь кандидата строго по профилю должности.

Правила:
1. Сначала оцени соответствие по категориям (profile_requirements_met), затем выведи общий score на их основе.
2. profile_requirements_met — это ЦЕЛЫЕ ПРОЦЕНТЫ от 0 до 100 (не доли 0.03, не баллы 0–4):
   - hard_skills: обязательные и желательные hard skills из профиля
   - soft_skills: психологические черты и коммуникативные качества
   - experience: опыт, анкетные требования, релевантность прошлых ролей
3. Согласованность score и процентов (ОБЯЗАТЕЛЬНО):
   - score 4 → каждая категория 90–100%, среднее ≥ 90%
   - score 3 → каждая категория 60–89%, среднее 70–85%
   - score 2 → каждая категория 30–79%, среднее 40–70%
   - score 1 → каждая категория 10–39%, среднее 15–35%
   - score 0 → каждая категория 0–29%, стоп-фактор или полный провал
4. В comment явно перечисли: какие требования профиля выполнены, частично, не выполнены; как категории повлияли на score.
5. strengths/weaknesses — ссылки на конкретные пункты профиля.
6. Если передан флаг игнорирования — не штрафуй за это требование и укажи в comment: «Требование X проигнорировано по флагу».
7. РАСШИФРОВКА СОБЕСЕДОВАНИЯ — это запись первичного интервью в разговорной форме:
   - Ответы кандидата часто косвенные, не по всем пунктам профиля; ищи подтверждения в примерах, деталях, формулировках.
   - Если требование не обсуждалось на интервью — укажи «не раскрыто на интервью», не ставь 0% автоматически; учти резюме.
   - Если в резюме есть навык, а на интервью не уточняли — оцени как частично подтверждённый (умеренный %), не как полный провал.
   - Сопоставляй фрагменты расшифровки с вопросами опросника (если передан) и требованиями профиля.
   - Не требуй дословных формулировок из профиля в ответах кандидата.
8. ОБЯЗАТЕЛЬНО проанализируй по расшифровке (и резюме, если есть):
   - Причины ухода с последнего места работы или перехода с фриланса/самозанятости; согласованность ответов с резюме.
   - Причины долгого поиска работы при пробелах в опыте.
   - Красные флаги: уход от прямого ответа, размытые формулировки, попытка скрыть истинную причину, односторонние обвинения работодателей, признаки конфликтности, излишней требовательности к условиям.
9. Оцени лояльность к условиям компании, адекватность и предсказуемость, управляемость — по ответам о мотивации, раздражителях на работе, готовности к обратной связи от прошлых работодателей. Отрази в soft_skills, comment, strengths/weaknesses.
10. Явные риски по п.8–9 снижают soft_skills и могут снизить итоговый score на 1 балл даже при сильных hard skills.
11. Если передан КОММЕНТАРИЙ HR — обязательно учти его: это замечания рекрутера после контакта с кандидатом; согласуй оценку и score с ними, отрази в comment.
12. Отвечай ТОЛЬКО валидным JSON без markdown."""


def normalize_category_percent(value):
    if value is None:
        return None
    try:
        v = float(value)
    except (TypeError, ValueError):
        return None
    if 0 < v <= 1:
        v *= 100
    return int(round(max(0, min(100, v))))


def normalize_evaluation_result(result):
    """Приводит score и profile_requirements_met к согласованным значениям."""
    try:
        score = int(float(result.get("score", 0)))
    except (TypeError, ValueError):
        score = 0
    score = max(0, min(4, score))
    result["score"] = score

    met = result.get("profile_requirements_met") or {}
    if not isinstance(met, dict):
        met = {}

    normalized = {key: normalize_category_percent(met.get(key)) for key in EVAL_CATEGORY_KEYS}
    low, high = SCORE_CATEGORY_RANGES[score]
    mid = (low + high) // 2

    present = [normalized[k] for k in EVAL_CATEGORY_KEYS if normalized[k] is not None]
    if present and max(present) <= 4 and score >= 2:
        normalized = {key: mid for key in EVAL_CATEGORY_KEYS}

    avg = sum(normalized[k] for k in EVAL_CATEGORY_KEYS if normalized[k] is not None) / len(EVAL_CATEGORY_KEYS) if present else 0
    if score >= 3 and avg < 20:
        if avg > 0:
            expected_mid = (low + high) / 2
            scale = expected_mid / avg
            for key in EVAL_CATEGORY_KEYS:
                if normalized[key] is not None:
                    normalized[key] = int(round(normalized[key] * scale))
        else:
            normalized = {key: mid for key in EVAL_CATEGORY_KEYS}

    for key in EVAL_CATEGORY_KEYS:
        if normalized[key] is None:
            normalized[key] = mid
        normalized[key] = max(low, min(high, normalized[key]))

    result["profile_requirements_met"] = normalized
    return result


def get_vacancy_questionnaire_text(vacancy):
    questions = vacancy.get("documents", {}).get("questions", "")
    if not questions or not str(questions).strip():
        return ""
    text = str(questions).strip()
    if text.startswith("["):
        try:
            return json.dumps(json.loads(text), ensure_ascii=False, indent=2)
        except json.JSONDecodeError:
            return text
    return text


QUESTIONNAIRE_JSON_SCHEMA = """{
  "опросник": [{
    "вопрос": "основной вопрос в разговорной форме",
    "уточняющие_вопросы": ["уточнение 1"],
    "проверяет_требование": "какой пункт профиля проверяем",
    "категория": "hard_skills | soft_skills | experience | motivation | reliability",
    "пример_ответа": "реалистичный ответ сильного кандидата"
  }]
}"""

QUESTIONNAIRE_REGENERATE_SYSTEM = f"""Ты — HR-директор с опытом проведения первичных собеседований.
Твоя задача — сформировать или пересобрать опросник для первичного интервью по профилю должности.

{QUESTIONNAIRE_GENERATION_RULES}

Если передан текущий опросник — улучши его с учётом коррективов, сохраняя удачные формулировки где уместно.
Верни ТОЛЬКО валидный JSON без markdown по схеме:
{QUESTIONNAIRE_JSON_SCHEMA}"""


def profile_to_text(profile):
    if isinstance(profile, dict):
        return json.dumps(profile, ensure_ascii=False, indent=2)
    return str(profile or "").strip()


def parse_questionnaire_input(questionnaire):
    if isinstance(questionnaire, list):
        return normalize_docs({"опросник": questionnaire})["опросник"]
    text = str(questionnaire or "").strip()
    if not text:
        return []
    if text.startswith("["):
        try:
            parsed = json.loads(text)
            if isinstance(parsed, list):
                return normalize_docs({"опросник": parsed})["опросник"]
        except json.JSONDecodeError:
            pass
    return normalize_docs({"опросник": [{"вопрос": text, "пример_ответа": ""}]})["опросник"]


def regenerate_questionnaire_with_ai(job_title, profile, current_questionnaire=None, corrections=""):
    profile_text = profile_to_text(profile)
    if not profile_text:
        raise ValueError("Профиль должности пуст — сначала сформируйте или загрузите профиль.")

    current_list = parse_questionnaire_input(current_questionnaire) if current_questionnaire else []
    user_parts = [f"Должность: {job_title or '—'}", f"ПРОФИЛЬ ДОЛЖНОСТИ:\n{profile_text}"]
    if current_list:
        current_json = json.dumps(current_list, ensure_ascii=False, indent=2)
        user_parts.append(f"ТЕКУЩИЙ ОПРОСНИК:\n{current_json}")
    if corrections and str(corrections).strip():
        user_parts.append(f"КОРРЕКТИВЫ ОТ HR (обязательно учти):\n{str(corrections).strip()}")
    user_parts.append("Сформируй опросник для первичного собеседования.")

    response = client.chat.completions.create(
        model=config["model"]["name"],
        messages=[
            {"role": "system", "content": QUESTIONNAIRE_REGENERATE_SYSTEM},
            {"role": "user", "content": "\n\n".join(user_parts)},
        ],
        temperature=config["model"]["temperature"],
        max_tokens=config["model"]["max_tokens"],
    )
    result = parse_ai_json_response(response.choices[0].message.content)
    raw = result.get("опросник", result if isinstance(result, list) else [])
    if not isinstance(raw, list) or not raw:
        raise ValueError("ИИ вернул пустой или некорректный опросник.")
    return normalize_docs({"опросник": raw})["опросник"]


PROFILE_REGENERATE_SYSTEM = """Ты — HR-директор. Обнови профиль должности по коррективам HR.
Верни ТОЛЬКО JSON: {"профиль": { структура как в образце профиля }}."""

VACANCY_TEXT_REGENERATE_SYSTEM = """Ты — HR-директор. Обнови текст вакансии по коррективам HR.
Верни ТОЛЬКО JSON: {"текст_вакансии": "..."}."""

KEYWORDS_REGENERATE_SYSTEM = """Ты — HR-рекрутер. Обнови ключевые слова для поиска кандидатов по коррективам HR.
Верни ТОЛЬКО JSON: {"ключевые_слова": ["слово1", "слово2"]}."""


def regenerate_profile_with_ai(job_title, profile, corrections=""):
    profile_text = profile_to_text(profile)
    if not profile_text:
        raise ValueError("Профиль пуст — нечего перегенерировать.")
    user_parts = [f"Должность: {job_title or '—'}", f"ТЕКУЩИЙ ПРОФИЛЬ:\n{profile_text}"]
    if corrections and str(corrections).strip():
        user_parts.append(f"КОРРЕКТИВЫ ОТ HR:\n{str(corrections).strip()}")
    response = client.chat.completions.create(
        model=config["model"]["name"],
        messages=[
            {"role": "system", "content": PROFILE_REGENERATE_SYSTEM},
            {"role": "user", "content": "\n\n".join(user_parts)},
        ],
        temperature=config["model"]["temperature"],
        max_tokens=config["model"]["max_tokens"],
    )
    result = parse_ai_json_response(response.choices[0].message.content)
    prof = result.get("профиль", result if isinstance(result, dict) and "задачи" in result else {})
    return normalize_docs({"профиль": prof})["профиль"]


def regenerate_vacancy_text_with_ai(job_title, profile, current_text, corrections=""):
    profile_text = profile_to_text(profile)
    user_parts = [
        f"Должность: {job_title or '—'}",
        f"ПРОФИЛЬ:\n{profile_text or '—'}",
        f"ТЕКУЩИЙ ТЕКСТ ВАКАНСИИ:\n{current_text or '—'}",
    ]
    if corrections and str(corrections).strip():
        user_parts.append(f"КОРРЕКТИВЫ ОТ HR:\n{str(corrections).strip()}")
    response = client.chat.completions.create(
        model=config["model"]["name"],
        messages=[
            {"role": "system", "content": VACANCY_TEXT_REGENERATE_SYSTEM},
            {"role": "user", "content": "\n\n".join(user_parts)},
        ],
        temperature=config["model"]["temperature"],
        max_tokens=config["model"]["max_tokens"],
    )
    result = parse_ai_json_response(response.choices[0].message.content)
    return result.get("текст_вакансии", current_text)


def regenerate_keywords_with_ai(job_title, profile, current_keywords, corrections=""):
    profile_text = profile_to_text(profile)
    kw_text = current_keywords if isinstance(current_keywords, str) else ", ".join(current_keywords or [])
    user_parts = [
        f"Должность: {job_title or '—'}",
        f"ПРОФИЛЬ:\n{profile_text or '—'}",
        f"ТЕКУЩИЕ КЛЮЧЕВЫЕ СЛОВА:\n{kw_text or '—'}",
    ]
    if corrections and str(corrections).strip():
        user_parts.append(f"КОРРЕКТИВЫ ОТ HR:\n{str(corrections).strip()}")
    response = client.chat.completions.create(
        model=config["model"]["name"],
        messages=[
            {"role": "system", "content": KEYWORDS_REGENERATE_SYSTEM},
            {"role": "user", "content": "\n\n".join(user_parts)},
        ],
        temperature=config["model"]["temperature"],
        max_tokens=800,
    )
    result = parse_ai_json_response(response.choices[0].message.content)
    kws = result.get("ключевые_слова", [])
    return ", ".join(kws) if isinstance(kws, list) else str(kws)


def render_questionnaire_item(index, q):
    if not isinstance(q, dict):
        st.markdown(f"**{index}. {q}**")
        return
    st.markdown(f"**{index}. {q.get('вопрос', '')}**")
    meta = []
    if q.get("проверяет_требование"):
        meta.append(f"Проверяет: {q['проверяет_требование']}")
    if q.get("категория"):
        meta.append(f"Категория: {q['категория']}")
    if meta:
        st.caption(" · ".join(meta))
    followups = q.get("уточняющие_вопросы", [])
    if followups:
        with st.expander("Уточняющие вопросы", expanded=False):
            for j, followup in enumerate(followups, 1):
                st.markdown(f"{j}. {followup}")
    if q.get("пример_ответа"):
        st.caption(f"Пример ответа: {q['пример_ответа']}")


def render_questionnaire_edit_panel(job_title, profile, questionnaire, key_prefix, on_apply):
    """Панель коррективов, перегенерации и ручного редактирования опросника."""
    questions = parse_questionnaire_input(questionnaire)
    st.markdown("##### ✏️ Коррективы и перегенерация")
    corrections = st.text_area(
        "Укажите, что изменить в опроснике",
        placeholder=(
            "Например: добавить вопрос про опыт с Excel; убрать вопрос про зарплату; "
            "сделать вопросы мягче; больше проверять навык «аналитика данных»"
        ),
        height=100,
        key=f"{key_prefix}_corrections",
    )
    if st.button("🔄 Перегенерировать опросник", key=f"{key_prefix}_regen", use_container_width=True):
        with st.spinner("Перегенерация опросника..."):
            try:
                new_questions = regenerate_questionnaire_with_ai(
                    job_title, profile, questions, corrections
                )
                on_apply(new_questions)
                st.success("Опросник перегенерирован!")
                st.rerun()
            except Exception as e:
                st.error(f"Ошибка перегенерации: {e}")

    with st.expander("📝 Редактировать опросник вручную (JSON)", expanded=False):
        json_value = json.dumps(questions, ensure_ascii=False, indent=2)
        edited_json = st.text_area(
            "JSON опросника",
            value=json_value,
            height=280,
            key=f"{key_prefix}_json_edit",
        )
        if st.button("✅ Сохранить ручные правки", key=f"{key_prefix}_save_json"):
            try:
                parsed = json.loads(edited_json)
                if not isinstance(parsed, list):
                    raise ValueError("Ожидается JSON-массив вопросов.")
                on_apply(normalize_docs({"опросник": parsed})["опросник"])
                st.success("Ручные правки сохранены!")
                st.rerun()
            except (json.JSONDecodeError, ValueError) as e:
                st.error(f"Некорректный JSON: {e}")

    st.divider()
    st.markdown("##### 📋 Текущий опросник")
    if not questions:
        st.info("Опросник пуст. Перегенерируйте его по профилю должности.")
    for i, q in enumerate(questions, 1):
        render_questionnaire_item(i, q)


def transcribe_speechkit_cloud(audio_path):
    """Расшифровка через Яндекс SpeechKit (файл загружается в Object Storage)."""
    _validate_speechkit_config(
        bucket=os.getenv("YANDEX_BUCKET_NAME"),
        access_key=os.getenv("YANDEX_ACCESS_KEY_ID"),
        secret_key=os.getenv("YANDEX_SECRET_ACCESS_KEY"),
        api_key=os.getenv("YANDEX_API_KEY"),
    )
    pcm_path = convert_to_pcm(audio_path)
    audio_url = upload_to_s3_and_get_url(
        pcm_path,
        os.getenv("YANDEX_BUCKET_NAME"),
        os.getenv("YANDEX_ACCESS_KEY_ID"),
        os.getenv("YANDEX_SECRET_ACCESS_KEY"),
    )
    return recognize_long_audio(audio_url, os.getenv("YANDEX_API_KEY"))


def generate_from_transcript(transcript_text, job_title="", doc_flags=None):
    """Генерирует пакет HR-документов из расшифровки (можно выбрать состав)."""
    flags = doc_flags or {
        "profile": True,
        "questionnaire": True,
        "vacancy_text": True,
        "keywords": True,
    }
    doc_parts = []
    if flags.get("profile"):
        doc_parts.append("профиль должности")
    if flags.get("questionnaire"):
        doc_parts.append("опросник для первичного собеседования")
    if flags.get("vacancy_text"):
        doc_parts.append("текст вакансии")
    if flags.get("keywords"):
        doc_parts.append("ключевые слова для поиска")
    docs_list = ", ".join(doc_parts) if doc_parts else "профиль должности и опросник"

    job_title_block = f'Поле "должность" в JSON: "{job_title}".' if job_title else ""
    system_prompt = f"""
    Ты — HR-директор с 15-летним опытом. На основе расшифровки создай ТОЛЬКО:
    {docs_list}.

    Образец профиля: {EXAMPLE_PROFILE}
    Примеры заполнения: {EXAMPLE_FILLED_PROFILE}
    Образец вакансии: {EXAMPLE_VACANCY}
    Образец опросника: {EXAMPLE_QUESTIONNAIRE}
    Пример опросника: {EXAMPLE_FILLED_QUESTIONNAIRE}
    {QUESTIONNAIRE_GENERATION_RULES}
    {job_title_block}

    Верни строго JSON с полями только для запрошенных документов:
    {{
      "должность": "...",
      "профиль": {{...}},
      "текст_вакансии": "...",
      "опросник": [{{"вопрос": "...", "уточняющие_вопросы": [], "проверяет_требование": "...", "категория": "...", "пример_ответа": "..."}}],
      "ключевые_слова": ["..."]
    }}
    Не включай в JSON поля документов, которые не запрашивались.
    """
    response = client.chat.completions.create(
        model=config["model"]["name"],
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"Текст расшифровки:\n{transcript_text}"},
        ],
        temperature=config["model"]["temperature"],
        max_tokens=config["model"]["max_tokens"],
    )
    result = parse_ai_json_response(response.choices[0].message.content)
    result = normalize_docs(result)
    if job_title:
        result["должность"] = job_title
    return result


def build_vacancy_deps():
    return {
        "client": client,
        "config": config,
        "load_vacancies": load_vacancies,
        "save_vacancies": save_vacancies,
        "load_chats": load_chats,
        "create_vacancy": create_vacancy,
        "create_vacancy_from_template": create_vacancy_from_template,
        "update_vacancy_docs": update_vacancy_docs,
        "update_vacancy_docs_by_id": update_vacancy_docs_by_id,
        "extract_text": extract_text,
        "extract_text_from_pdf_url": extract_text_from_pdf_url,
        "get_direct_yandex_link": get_direct_yandex_link,
        "normalize_docs": normalize_docs,
        "parse_ai_json_response": parse_ai_json_response,
        "regenerate_questionnaire_with_ai": regenerate_questionnaire_with_ai,
        "render_questionnaire_edit_panel": render_questionnaire_edit_panel,
        "save_generation_to_history": save_generation_to_history,
        "export_to_word": export_to_word,
        "export_to_pdf": export_to_pdf,
        "generate_from_transcript": generate_from_transcript,
        "transcribe_speechkit_cloud": transcribe_speechkit_cloud,
        "transcribe_video_from_link": transcribe_video_from_link,
        "regenerate_profile_with_ai": regenerate_profile_with_ai,
        "regenerate_vacancy_text_with_ai": regenerate_vacancy_text_with_ai,
        "regenerate_keywords_with_ai": regenerate_keywords_with_ai,
        "profile_to_text": profile_to_text,
        "parse_questionnaire_input": parse_questionnaire_input,
        "render_questionnaire_item": render_questionnaire_item,
        "vacancy_has_profile": vacancy_has_profile,
        "get_vacancy_profile_text": get_vacancy_profile_text,
        "get_vacancy_questionnaire_text": get_vacancy_questionnaire_text,
        "evaluate_candidate_with_ai_v2": evaluate_candidate_with_ai_v2,
        "get_history_index": get_history_index,
        "load_generation_from_history": load_generation_from_history,
        "delete_generation_from_history": delete_generation_from_history,
        "send_telegram_message": send_telegram_message,
        "default_ignore_flags": default_ignore_flags,
        "QUESTIONNAIRE_GENERATION_RULES": QUESTIONNAIRE_GENERATION_RULES,
        "EXAMPLE_PROFILE": EXAMPLE_PROFILE,
        "EXAMPLE_FILLED_PROFILE": EXAMPLE_FILLED_PROFILE,
        "EXAMPLE_VACANCY": EXAMPLE_VACANCY,
        "EXAMPLE_QUESTIONNAIRE": EXAMPLE_QUESTIONNAIRE,
        "EXAMPLE_FILLED_QUESTIONNAIRE": EXAMPLE_FILLED_QUESTIONNAIRE,
    }


def evaluate_candidate_with_ai_v2(
    resume_text,
    transcript_text,
    job_title,
    vacancy_profile,
    ignore_flags,
    interview_questionnaire="",
    interview_eval_notes="",
):
    from ai_helpers import create_chat_completion, get_char_limit, trim_profile_for_eval, trim_text
    from resume_ai import format_interview_eval_notes_block, questionnaire_to_eval_prompt

    profile_block = trim_profile_for_eval(
        vacancy_profile,
        get_char_limit(config, "profile", 5000),
    )
    resume_block = trim_text(resume_text or "Не предоставлено", get_char_limit(config, "resume", 8000))
    transcript_block = trim_text(
        transcript_text
        if transcript_text
        else "Не предоставлена — опирайся на резюме, снижай уверенность оценки",
        get_char_limit(config, "transcript", 10000),
    )
    questionnaire_trimmed = trim_text(
        questionnaire_to_eval_prompt(interview_questionnaire) or interview_questionnaire or "",
        get_char_limit(config, "questionnaire", 4000),
    )

    questionnaire_block = ""
    if questionnaire_trimmed.strip():
        questionnaire_block = f"""
ОПРОСНИК ПЕРВИЧНОГО СОБЕСЕДОВАНИЯ ДЛЯ ЭТОГО КАНДИДАТА (вопросы, эталоны ответов и оценки HR по каждому пункту):
{questionnaire_trimmed}

Сопоставь расшифровку интервью с КАЖДЫМ вопросом опросника.
Учитывай оценку HR по каждому вопросу как мнение рекрутера, присутствовавшего на интервью:
- «Хорошо» — HR полностью удовлетворён; не занижай без веских причин в расшифровке.
- «Удовлетворительно» — в целом ок с небольшими недочётами.
- «Сомнительно» — HR не увидел однозначного соответствия или заметил тревожный фактор; можешь согласиться или мягко оспорить со ссылкой на расшифровку.
- «Нет» — ответ неудовлетворителен.
Если твоё мнение расходится с оценкой HR — явно укажи это в comment.
"""

    user_prompt = f"""Оцени кандидата на позицию: {job_title}

ПРОФИЛЬ ДОЛЖНОСТИ:
{profile_block}

{questionnaire_block}

{build_ignore_flags_prompt(ignore_flags)}

РЕЗЮМЕ КАНДИДАТА:
{resume_block}

РАСШИФРОВКА ПЕРВИЧНОГО СОБЕСЕДОВАНИЯ (разговорная речь, ответы могут быть неполными):
{transcript_block}
{format_interview_eval_notes_block(interview_eval_notes, config)}
Алгоритм:
1. Составь чек-лист требований профиля; для каждого отметь: подтверждено интервью / подтверждено резюме / частично / не раскрыто.
2. Отдельно разбери: причины ухода/поиска работы, пробелы в опыте, мотивацию, раздражители, обратную связь от работодателей — лояльность, адекватность, управляемость, конфликтность, требовательность.
3. Разбери соответствие по категориям hard_skills / soft_skills / experience (проценты 0–100); soft_skills включает п.2.
4. На основе процентов, стоп-факторов, оценок HR по вопросам опросника и рисков п.2 определи score 0–4.
5. В comment: по каждой категории — что услышано на интервью; отдельный абзац про мотивацию/уход/управляемость; если есть уточнения HR для оценки — как они повлияли; при расхождении с оценкой HR по вопросу — поясни почему.

Верни JSON:
{{
  "score": 3,
  "comment": "Hard skills: 75% — выполнено X, не хватает Y. Soft skills: 70% — ... Experience: 80% — ... Итоговый score 3/4, потому что ...",
  "strengths": ["соответствует требованию X из профиля"],
  "weaknesses": ["отсутствует требование Y из профиля"],
  "profile_requirements_met": {{
    "hard_skills": 75,
    "soft_skills": 70,
    "experience": 80
  }},
  "flags_applied": []
}}

ВАЖНО: profile_requirements_met — только целые проценты 0–100. Не используй 0.75 вместо 75 и не ставь 3–4 как процент при score 3."""
    try:
        response = create_chat_completion(
            client,
            config,
            "interview_eval",
            messages=[
                {"role": "system", "content": EVAL_SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.3,
        )
        result = parse_ai_json_response(response.choices[0].message.content)
        normalized = normalize_evaluation_result(result)
        normalized["ok"] = True
        return normalized
    except Exception as e:
        return {
            "ok": False,
            "score": None,
            "error": str(e),
            "comment": f"Ошибка оценки: {e}",
            "strengths": [],
            "weaknesses": [],
            "profile_requirements_met": {},
            "flags_applied": [],
        }


def evaluate_candidate_with_ai(resume_text, transcript_text, job_title):
    """Устаревшая оценка без профиля. Используйте evaluate_candidate_with_ai_v2."""
    return evaluate_candidate_with_ai_v2(
        resume_text, transcript_text, job_title, "Профиль не указан.", default_ignore_flags()
    )

# ============================================================
# НОРМАЛИЗАЦИЯ JSON
# ============================================================
def normalize_docs(doc):
    if not isinstance(doc, dict):
        return doc
    if 'профиль' in doc and isinstance(doc['профиль'], dict):
        profile = doc['профиль']
        for field in ['обязательные_требования', 'желательные_требования', 'психологические_черты']:
            if field in profile and isinstance(profile[field], list):
                new_items = []
                for item in profile[field]:
                    if isinstance(item, str):
                        if field == 'психологические_черты':
                            new_items.append({"качество": item, "проявление": item})
                        else:
                            new_items.append({"навык": item, "описание": item})
                    elif isinstance(item, dict):
                        new_items.append(item)
                    else:
                        if field == 'психологические_черты':
                            new_items.append({"качество": str(item), "проявление": str(item)})
                        else:
                            new_items.append({"навык": str(item), "описание": str(item)})
                profile[field] = new_items
        if 'задачи' in profile and isinstance(profile['задачи'], list):
            profile['задачи'] = [str(t) for t in profile['задачи']]
    if 'опросник' in doc and isinstance(doc['опросник'], list):
        new_q = []
        for q in doc['опросник']:
            if isinstance(q, str):
                new_q.append({"вопрос": q, "пример_ответа": ""})
            elif isinstance(q, dict):
                followups = q.get("уточняющие_вопросы", q.get("followups", []))
                if isinstance(followups, str):
                    followups = [followups] if followups.strip() else []
                new_q.append({
                    "вопрос": q.get("вопрос", q.get("question", str(q))),
                    "уточняющие_вопросы": [str(f) for f in followups] if isinstance(followups, list) else [],
                    "проверяет_требование": q.get("проверяет_требование", q.get("requirement", "")),
                    "категория": q.get("категория", q.get("category", "")),
                    "пример_ответа": q.get("пример_ответа", q.get("example", "")),
                })
            else:
                new_q.append({"вопрос": str(q), "пример_ответа": ""})
        doc['опросник'] = new_q
    if 'ключевые_слова' in doc and isinstance(doc['ключевые_слова'], list):
        doc['ключевые_слова'] = [str(kw) for kw in doc['ключевые_слова']]
    return doc

# ============================================================
# ЭКСПОРТ В WORD И PDF
# ============================================================
def export_to_word(gen_data):
    doc = Document()
    title = doc.add_heading(f"{config['app']['name']} v{config['app']['version']}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading("ПРОФИЛЬ ДОЛЖНОСТИ", level=1)
    profile = gen_data.get("профиль", {})
    doc.add_heading(f"Должность: {gen_data.get('должность', '—')}", level=2)
    doc.add_paragraph(f"Подразделение: {profile.get('подразделение', '—')}")
    doc.add_paragraph(f"Непосредственный руководитель: {profile.get('непосредственный_руководитель', '—')}")
    
    doc.add_heading("1. Задачи", level=2)
    tasks = profile.get("задачи", [])
    for task in tasks:
        doc.add_paragraph(task, style='List Bullet')
    
    doc.add_heading("2. Анкетные требования", level=2)
    at = profile.get("анкетные_требования", {})
    doc.add_paragraph(f"Возраст: {at.get('возраст', '—')}")
    doc.add_paragraph(f"Пол: {at.get('пол', '—')}")
    doc.add_paragraph("Стоп-факторы:")
    for sf in at.get("стоп_факторы", []):
        doc.add_paragraph(sf, style='List Bullet')
    
    doc.add_heading("3. Обязательные требования (Hard Skills)", level=2)
    for req in profile.get("обязательные_требования", []):
        if isinstance(req, dict):
            doc.add_paragraph(f"{req.get('навык', '')}: {req.get('описание', '')}", style='List Bullet')
        else:
            doc.add_paragraph(str(req), style='List Bullet')
    
    doc.add_heading("4. Желательные требования", level=2)
    for req in profile.get("желательные_требования", []):
        if isinstance(req, dict):
            doc.add_paragraph(f"{req.get('навык', '')}: {req.get('описание', '')}", style='List Bullet')
        else:
            doc.add_paragraph(str(req), style='List Bullet')
    
    doc.add_heading("5. Психологические черты (Soft Skills)", level=2)
    for trait in profile.get("психологические_черты", []):
        if isinstance(trait, dict):
            doc.add_paragraph(f"{trait.get('качество', '')}: {trait.get('проявление', '')}", style='List Bullet')
        else:
            doc.add_paragraph(str(trait), style='List Bullet')
    
    doc.add_heading("6. Условия работы", level=2)
    cond = profile.get("условия_работы", {})
    doc.add_paragraph(f"Формат: {cond.get('формат', '—')}")
    doc.add_paragraph(f"Режим: {cond.get('режим', '—')}")
    doc.add_paragraph(f"Зарплата: {cond.get('зарплата', '—')}")
    doc.add_paragraph(f"Испытательный срок: {cond.get('испытательный_срок', '—')}")
    
    doc.add_page_break()
    doc.add_heading("ТЕКСТ ВАКАНСИИ", level=1)
    doc.add_paragraph(gen_data.get("текст_вакансии", ""))
    
    doc.add_page_break()
    doc.add_heading("ОПРОСНИК ДЛЯ ПЕРВИЧНОГО СОБЕСЕДОВАНИЯ", level=1)
    for i, q in enumerate(gen_data.get("опросник", []), 1):
        doc.add_heading(f"{i}. {q.get('вопрос', '')}", level=2)
        if q.get("проверяет_требование"):
            doc.add_paragraph(f"Проверяет: {q.get('проверяет_требование')}")
        if q.get("категория"):
            doc.add_paragraph(f"Категория: {q.get('категория')}")
        for j, followup in enumerate(q.get("уточняющие_вопросы", []), 1):
            doc.add_paragraph(f"Уточняющий {j}: {followup}")
        doc.add_paragraph(f"Пример желательного ответа: {q.get('пример_ответа', '')}", style='Intense Quote')
    
    doc.add_heading("КЛЮЧЕВЫЕ СЛОВА ДЛЯ ПОИСКА", level=1)
    keywords = gen_data.get("ключевые_слова", [])
    doc.add_paragraph(", ".join(keywords))
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        return tmp.name

def export_to_pdf(gen_data):
    font_dir = "fonts"
    os.makedirs(font_dir, exist_ok=True)
    font_path = os.path.join(font_dir, "DejaVuSans.ttf")
    
    if not os.path.exists(font_path):
        st.error(
            "❌ Не найден шрифт для генерации PDF.\n\n"
            "Скачайте файл DejaVuSans.ttf и поместите его в папку 'fonts' в корне проекта."
        )
        return None

    pdfmetrics.registerFont(TTFont('DejaVu', font_path))
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        pdf_path = tmp.name

    c = canvas.Canvas(pdf_path, pagesize=A4)
    width, height = A4
    y = height - 20
    line_height = 14

    def draw_text(text, font_size=11, bold=False):
        nonlocal y
        if y < 50:
            c.showPage()
            y = height - 20
        c.setFont('DejaVu', font_size)
        lines = text.split('\n')
        for line in lines:
            if y < 50:
                c.showPage()
                y = height - 20
            c.drawString(20, y, line)
            y -= line_height

    def draw_bullet_list(items):
        nonlocal y
        for item in items:
            if y < 50:
                c.showPage()
                y = height - 20
            c.setFont('DejaVu', 11)
            c.drawString(25, y, f"• {item}")
            y -= line_height

    def draw_section_title(title):
        nonlocal y
        if y < 60:
            c.showPage()
            y = height - 20
        c.setFont('DejaVu', 14)
        c.drawString(20, y, title)
        y -= line_height + 5

    draw_section_title("ПРОФИЛЬ ДОЛЖНОСТИ")
    profile = gen_data.get("профиль", {})
    draw_text(f"Должность: {gen_data.get('должность', '—')}", font_size=12)
    draw_text(f"Подразделение: {profile.get('подразделение', '—')}")
    draw_text(f"Руководитель: {profile.get('непосредственный_руководитель', '—')}")

    draw_section_title("1. Задачи")
    draw_bullet_list(profile.get("задачи", []))

    draw_section_title("2. Анкетные требования")
    at = profile.get("анкетные_требования", {})
    draw_text(f"Возраст: {at.get('возраст', '—')}")
    draw_text(f"Пол: {at.get('пол', '—')}")
    draw_text("Стоп-факторы:")
    draw_bullet_list(at.get("стоп_факторы", []))

    draw_section_title("3. Обязательные требования")
    for req in profile.get("обязательные_требования", []):
        if isinstance(req, dict):
            draw_text(f"• {req.get('навык', '')}: {req.get('описание', '')}")
        else:
            draw_text(f"• {req}")

    draw_section_title("4. Желательные требования")
    for req in profile.get("желательные_требования", []):
        if isinstance(req, dict):
            draw_text(f"• {req.get('навык', '')}: {req.get('описание', '')}")
        else:
            draw_text(f"• {req}")

    draw_section_title("5. Психологические черты")
    for trait in profile.get("психологические_черты", []):
        if isinstance(trait, dict):
            draw_text(f"• {trait.get('качество', '')}: {trait.get('проявление', '')}")
        else:
            draw_text(f"• {trait}")

    draw_section_title("6. Условия работы")
    cond = profile.get("условия_работы", {})
    draw_text(f"Формат: {cond.get('формат', '—')}")
    draw_text(f"Режим: {cond.get('режим', '—')}")
    draw_text(f"Зарплата: {cond.get('зарплата', '—')}")
    draw_text(f"Испытательный срок: {cond.get('испытательный_срок', '—')}")

    draw_section_title("ТЕКСТ ВАКАНСИИ")
    draw_text(gen_data.get("текст_вакансии", ""))

    draw_section_title("ОПРОСНИК ДЛЯ СОБЕСЕДОВАНИЯ")
    for i, q in enumerate(gen_data.get("опросник", []), 1):
        draw_text(f"{i}. {q.get('вопрос', '')}", font_size=11)
        if q.get("проверяет_требование"):
            draw_text(f"   Проверяет: {q.get('проверяет_требование')}", font_size=9)
        if q.get("категория"):
            draw_text(f"   Категория: {q.get('категория')}", font_size=9)
        for j, followup in enumerate(q.get("уточняющие_вопросы", []), 1):
            draw_text(f"   Уточняющий {j}: {followup}", font_size=9)
        draw_text(f"   Пример: {q.get('пример_ответа', '')}", font_size=10)
        y -= 5

    draw_section_title("КЛЮЧЕВЫЕ СЛОВА")
    draw_text(", ".join(gen_data.get("ключевые_слова", [])))

    c.save()
    return pdf_path

# ============================================================
# ИНТЕРФЕЙС STREAMLIT
# ============================================================
st.set_page_config(
    page_title=f"{config['app']['name']} v{config['app']['version']}",
    page_icon="🧠",
    layout="wide",
)
apply_corporate_ui()

st.title(f"🧠 {config['app']['name']} v{config['app']['version']}")
st.markdown("**Разработчик:** А.А. Крупин")
st.caption("HR-платформа: подготовка вакансий, воронка кандидатов, оценка ИИ и клиентская зона.")

# Боковая панель
departments = load_departments()
with st.sidebar:
    st.markdown('<p class="sidebar-section-label">Клиентские зоны</p>', unsafe_allow_html=True)
    st.markdown(
        '<a class="client-zone-btn" href="/master" target="_self">🏢 Мастер-зона (руководитель)</a>',
        unsafe_allow_html=True,
    )
    client_zone_links = "".join(
        f'<a class="client-zone-btn" href="/client?dept={quote(dept["name"])}" target="_self">'
        f'{"🧪 " if dept.get("id") == 99 or dept.get("slug") == "test" else ""}{dept["name"]}</a>'
        for dept in departments
    )
    st.markdown(client_zone_links, unsafe_allow_html=True)

    st.divider()
    st.markdown('<p class="sidebar-section-label">Проверка баланса</p>', unsafe_allow_html=True)
    st.markdown(
        """
        <div class="sidebar-links-group">
            <a class="sidebar-link-btn" href="https://console.yandex.cloud/folders/b1glrlal8l5f9uu25jjr/dashboard" target="_blank" rel="noopener noreferrer">Яндекс Облако</a>
            <a class="sidebar-link-btn" href="https://routerai.ru/settings/billing" target="_blank" rel="noopener noreferrer">RouterAI</a>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.divider()
    st.markdown(
        f"""
        <div class="sidebar-brand">
            <div class="sidebar-brand-title">{config['app']['name']}</div>
            <div class="sidebar-brand-subtitle">v{config['app']['version']} · HR-платформа</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown('<p class="sidebar-section-label">Конфигурация</p>', unsafe_allow_html=True)
    st.markdown(
        f"""
        <div class="sidebar-config-card">
            <div class="sidebar-config-row">
                <span class="label">Модель</span>
                <span class="value">{config['model']['name']}</span>
            </div>
            <div class="sidebar-config-row">
                <span class="label">Температура</span>
                <span class="value">{config['model']['temperature']}</span>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    if st.button("Перезагрузить конфиг"):
        config = load_config()
        st.rerun()


# Вкладки
tab_vacancies, tab_stats, tab_settings, tab5 = st.tabs([
    "🏢 Вакансии",
    "📊 Статистика",
    "⚙️ Настройки",
    "📖 Инструкции",
])



with tab_vacancies:
    render_vacancy_tab(build_vacancy_deps())

with tab_stats:
    render_stats_tab(build_vacancy_deps(), create_vacancy_fn=create_vacancy)

with tab5:
    st.header("Инструкции по работе с HR-помогатором")
    st.markdown("""
Приложение помогает готовить документы по вакансии, вести кандидатов и
согласовывать решения с руководителем через Telegram. ИИ и расшифровка аудио
работают в облаке (RouterAI и Яндекс SpeechKit), на компьютере ничего тяжёлого
устанавливать не нужно.

**Вкладки приложения**

**Вакансии** — основная работа. Внутри: «Вакансии в работе», **«Архив»**, **«Поиск»** (кандидаты по ФИО и резюме),
«Шаблоны» и «Создание новой вакансии». Откройте вакансию кнопкой в списке —
увидите кандидатов, документы и статистику. Закрытую вакансию можно перенести
в архив. Удаление — отдельным блоком внизу карточки, в два шага, с подтверждением.

Одно и то же название должности можно использовать снова, если предыдущая
итерация уже в архиве. Две активные вакансии с одинаковым названием создать
нельзя — это разные записи с разными датами и кандидатами, но в работе одновременно
должна быть только одна.

**Шаблоны** — сохранённые пакеты документов и привязка к чату. Удобно, когда
должность повторяется: создаёте вакансию из шаблона, правите детали, не начиная
с нуля. Шаблон можно обновить из текущей вакансии или при закрытии в архив.

**Создание новой вакансии** — регистрация вакансии (название и чат Telegram),
затем генерация или импорт документов. Расшифровка разговора с заказчиком —
через загрузку аудио или видео (нужен установленный ffmpeg) или готовый текст.

**Прошлые генерации** — сохраняются автоматически в `data/history`. Подставить пакет
можно в «Документы по вакансии» → «Прошлые генерации» у открытой вакансии; при создании
вакансии с похожим названием показывается подсказка.

**Статистика** — продуктивность за месяц/квартал/полугодие, сравнение с прошлым периодом
и ИИ-анализ по кнопке; реестр «На гарантии» и детальная статистика по вакансии.

**Настройки** — чаты Telegram по отделам, список всех вакансий, шаблоны,
подключение Google Calendar. Удаление вакансии из общего списка — тоже здесь,
с подтверждением.

**Инструкции** (эта вкладка) — краткое руководство по работе с приложением,
Telegram-ботом и клиентскими зонами.

**Боковая панель** — ссылки на клиентские зоны (`/master` и `/client?dept=...`)
для руководителей, проверка баланса облачных сервисов.
    """)

    with st.expander("Вакансии и кандидаты"):
        st.markdown("""
В карточке кандидата — этапы воронки, оценка ИИ по резюме и интервью, отправка
в Telegram-чат, назначение собеседования. При смене этапа с «Назначено собеседование»
можно включить «Не удалять событие из Google Calendar», если встреча в календаре
должна остаться, а этап в приложении меняется.

Кнопка «Отправить в общий чат» публикует карточку руководителю. Бот должен быть
запущен отдельно (`python bot.py` в терминале). Chat ID группы — в Настройках
или командой `/chatid` в группе.
        """)

    with st.expander("Работа в Telegram-чате"):
        st.markdown("""
Когда HR отправляет кандидата в чат, под сообщением появляется карточка с кнопками.
Это основной способ принять решение по кандидату.

**Что означают статусы**

«Встреча» — кандидата берут на собеседование. Статус сохраняется сразу. Дальше
нужно выбрать дату, время и формат: в офисе, удалённо или оба варианта. После
назначения HR получает запрос подтвердить встречу.

«Оффер» — готовы сделать предложение о работе. Сохраняется сразу, комментарий
не обязателен.

«Подумать» — нужно время на решение. Бот попросит написать комментарий ответом
на его сообщение. Пока комментарий не отправлен, статус не изменится. Пустое
сообщение бот не примет — напишите хотя бы короткое пояснение, почему откладываете
решение.

«Отказ» — кандидат не подходит. Тоже нужен комментарий: причина отказа важна
для HR и для истории по вакансии. Без комментария статус не сохранится.

**Комментарий отдельно от статуса**

Кнопка «Комментарий» добавляет пояснение, не меняя статус. Удобно, когда решение
уже принято, а нужно дописать детали. Можно также ответить текстом на карточку
кандидата — бот воспримет это как комментарий.

**Если передумали**

«Сменить статус» снова показывает все кнопки статусов. «Отменить встречу» сбрасывает
дату и время, если встреча не состоится.

**Напоминания**

Бот сам напомнит, если кандидат больше суток ждёт оценки, если статус «Подумать»
держится больше пяти дней, и за час до назначенной встречи (не раньше 08:00 по Москве).
По вторникам в 18:00 и по пятницам в 15:00 (по Москве) приходит сводка по кандидатам.
В субботу и воскресенье автоматические напоминания не уходят; накопленные за выходные
отправляются в понедельник в 10:00 по Москве.

**Команды в чате**

`/meetings` — предстоящие собеседования (подтверждены HR, ещё не прошли).
`/candidates` — просмотр карточек кандидатов по вакансии, можно указать название.
`/pending` — список тех, кто ещё ждёт оценки. Команды видны в меню слева от поля
ввода, если бот запущен и добавлен в группу.

**На что обратить внимание**

Кнопки статусов есть только у новых карточек, отправленных после запуска бота.
Если карточка старая — отправьте кандидата заново из приложения. Бот должен
быть запущен на вашем компьютере (`python bot.py`); одновременно не держите
два экземпляра с одним токеном.
        """)

    with st.expander("Клиентские зоны и доступ"):
        st.markdown("""
**Мастер-зона** (`/master`) — сводка по всем вакансиям для руководителя компании.

**Зона отдела** (`/client?dept=Название`) — кандидаты одного подразделения.
Название в ссылке должно совпадать с отделом в настройках.

Главную страницу приложения (`/`) руководителям лучше не давать — только эти ссылки.
При локальной работе ссылки работают, пока запущен Streamlit на вашем компьютере
(обычно `http://localhost:8501/...`). Для доступа из интернета позже можно
вынести приложение на сервер и указать `PUBLIC_APP_BASE_URL` в `.env`.
        """)

    with st.expander("Запуск бота и технические мелочи"):
        st.markdown("""
Бот запускается отдельно от сайта:

```bash
./venv/bin/python bot.py
```

В `.env` должен быть `TELEGRAM_BOT_TOKEN`. Для подтверждения встреч HR —
`TELEGRAM_HR_CONFIRM_USERNAME`. Для расшифровки — ключи Яндекс SpeechKit и ffmpeg.

Свой Telegram user id: напишите боту в личку `/id`.
        """)

# ---------- ВКЛАДКА: НАСТРОЙКИ ----------
with tab_settings:
    st.header("⚙️ Настройки")
    st.caption("Telegram-чаты, список вакансий и служебная информация.")

    with st.expander("📂 Мои чаты Telegram", expanded=True):
        chats = load_chats()
        departments = load_departments()

        col1, col2, col3, col4 = st.columns([2, 2, 2, 1])
        with col1:
            new_chat_name = st.text_input("Название чата", key="settings_chat_name")
        with col2:
            new_chat_id = st.text_input("Chat ID", key="settings_chat_id")
        with col3:
            dept_options = [d["name"] for d in departments]
            dept_choice = st.selectbox(
                "Подразделение",
                dept_options + ["➕ Создать новое..."],
                help="Выберите подразделение или создайте новое",
                key="settings_dept_choice",
            )
            new_dept_name = None
            if dept_choice == "➕ Создать новое...":
                new_dept_name = st.text_input("Название нового подразделения", key="settings_new_dept")
        with col4:
            st.write("")

        if st.button("💾 Сохранить чат", key="settings_save_chat"):
            if new_chat_name and new_chat_id:
                if dept_choice == "➕ Создать новое...":
                    if new_dept_name:
                        dept_id = add_department(new_dept_name)
                        dept_name = new_dept_name
                    else:
                        st.error("Введите название нового подразделения")
                        st.stop()
                else:
                    dept = next(d for d in departments if d["name"] == dept_choice)
                    dept_id = dept["id"]
                    dept_name = dept["name"]

                if not any(c["name"] == new_chat_name for c in chats):
                    from telegram_notify import normalize_chat_id
                    normalized_id = normalize_chat_id(new_chat_id)
                    if normalized_id is None:
                        st.error("Некорректный Chat ID")
                        st.stop()
                    chats.append({
                        "name": new_chat_name,
                        "id": normalized_id,
                        "department_name": dept_name,
                        "department_id": dept_id,
                    })
                    save_chats(chats)
                    from vacancy_store import (
                        prune_stale_telegram_posts,
                        sync_vacancy_chat_ids_from_chats,
                    )
                    if sync_vacancy_chat_ids_from_chats():
                        st.caption("Chat ID вакансий отдела обновлён.")
                    prune_stale_telegram_posts()
                    st.success("Сохранено!")
                    st.rerun()
                else:
                    st.error("Чат с таким именем уже есть.")
            else:
                st.warning("Заполните название чата и Chat ID.")

        if chats:
            st.write("Сохранённые чаты:")
            for i, chat in enumerate(chats):
                col_a, col_b, col_c, col_d = st.columns([2, 2, 2, 1])
                col_a.write(f"**{chat['name']}**")
                col_b.code(chat["id"])
                col_c.write(f"_{chat.get('department_name', '—')}_")
                if col_d.button("🗑️", key=f"settings_del_chat_{i}"):
                    chats.pop(i)
                    save_chats(chats)
                    from vacancy_store import sync_vacancy_chat_ids_from_chats
                    sync_vacancy_chat_ids_from_chats()
                    st.rerun()
        else:
            st.info("Чаты не добавлены. Сохраните первый чат выше.")

    with st.expander("🛡️ Гарантия", expanded=False):
        from app_settings import get_default_warranty_months, set_default_warranty_months
        from warranty import WARRANTY_MONTH_CHOICES, WARRANTY_MONTH_LABELS

        st.caption(
            "Срок по умолчанию подставляется при первом указании гарантии у кандидата. "
            "Месяц = 30 дней."
        )
        current_months = get_default_warranty_months()
        picked_months = st.selectbox(
            "Срок гарантии по умолчанию",
            WARRANTY_MONTH_CHOICES,
            index=WARRANTY_MONTH_CHOICES.index(current_months),
            format_func=lambda m: WARRANTY_MONTH_LABELS.get(m, f"{m} мес"),
            key="settings_default_warranty_months",
        )
        if st.button("💾 Сохранить срок гарантии", key="settings_save_warranty_default"):
            set_default_warranty_months(picked_months)
            st.success(f"Сохранено: {WARRANTY_MONTH_LABELS.get(picked_months, picked_months)}")

    with st.expander("🤖 Telegram-бот", expanded=True):
        from telegram_notify import get_bot_status, get_hr_user_id, send_telegram_html

        ok, status_msg, bot_info = get_bot_status()
        if ok:
            st.success(status_msg)
            st.caption(f"Имя: {bot_info.get('first_name', '—')}")
        else:
            st.error(status_msg)

        hr_id = get_hr_user_id()
        if hr_id:
            st.caption(f"HR user_id: `{hr_id}`")

        st.markdown(
            """
**Чтобы бот отвечал на команды** (`/id`, `/start`, `/chatid`) **и обрабатывал кнопки в чате**, запустите его в отдельном терминале:

```bash
./venv/bin/python bot.py
```

Процесс должен работать постоянно (рядом со Streamlit).

**Как узнать Chat ID группы:** добавьте бота в группу → напишите `/chatid` или `/id`.

**Как узнать свой user_id:** напишите боту в личку `/id`.

**Отправка кандидатов из приложения** работает через API (бот может быть выключен), но бот должен быть **добавлен в группу** вакансии.
            """
        )

        test_chats = load_chats()
        if ok and test_chats:
            test_chat_names = [c["name"] for c in test_chats]
            test_pick = st.selectbox(
                "Проверить отправку в чат",
                [""] + test_chat_names,
                key="settings_tg_test_chat",
            )
            if test_pick and st.button("📨 Отправить тестовое сообщение", key="settings_tg_test_btn"):
                chat = next(c for c in test_chats if c["name"] == test_pick)
                t_ok, t_msg, _ = send_telegram_html(
                    chat["id"],
                    "<b>✅ Тест</b>\n\nHR-помогатор успешно отправляет сообщения в этот чат.",
                )
                if t_ok:
                    st.success(t_msg)
                else:
                    st.error(t_msg)

    with st.expander("📅 Google Calendar", expanded=False):
        try:
            from google_calendar import (
                get_calendar_status,
                get_credentials_path,
                get_calendar_id,
                get_event_duration_minutes,
                credentials_file_exists,
            )
        except ImportError:
            st.warning(
                "Установите зависимости: `pip install google-api-python-client google-auth-oauthlib google-auth-httplib2`"
            )
        else:
            status, status_msg = get_calendar_status()
            if status == "ready":
                st.success(status_msg)
            elif status == "needs_auth":
                st.info(status_msg)
            else:
                st.warning(status_msg)

            st.markdown(
                """
**Настройка (один раз):**
1. [Google Cloud Console](https://console.cloud.google.com/) → APIs & Services → **Enable Google Calendar API**
2. Credentials → **OAuth client ID** → тип **Desktop app** → скачайте JSON
3. Сохраните файл как `data/google_calendar_credentials.json`
4. Нажмите **Подключить Google Calendar** — откроется браузер для входа
5. В `.env` при необходимости: `GOOGLE_CALENDAR_ID=primary` (или ID календаря), `GOOGLE_CALENDAR_EVENT_MINUTES=30`
                """
            )
            st.caption(
                f"Credentials: `{get_credentials_path()}` · Календарь: `{get_calendar_id()}` · "
                f"Длительность события: **{get_event_duration_minutes()} мин**"
            )

            if credentials_file_exists():
                st.caption(
                    "Если кнопка не открывает браузер — в Терминале: "
                    "`python google_calendar_auth.py` (из папки проекта с активным venv)."
                )
                if st.button("🔗 Подключить Google Calendar", key="settings_gcal_auth"):
                    from google_calendar import run_oauth_authorization

                    with st.spinner("Откроется браузер для авторизации…"):
                        ok, msg = run_oauth_authorization()
                    if ok:
                        st.success(msg)
                        st.rerun()
                    else:
                        st.error(msg)

            st.caption(
                "При статусе «Собеседование назначено» с датой и временем создаётся событие "
                "(например: «Иванов Иван, графический дизайнер»)."
            )

    st.divider()
    st.subheader("📌 Шаблоны вакансий")
    st.caption(
        "Полное редактирование документов — во вкладке «Вакансии» → «Шаблоны». "
        "Здесь можно быстро удалить шаблон."
    )
    from vacancy_template_store import delete_template, list_templates

    templates = list_templates()
    if not templates:
        st.info("Шаблонов пока нет.")
    else:
        for i, tpl in enumerate(templates):
            updated = (tpl.get("updated_at") or tpl.get("created_at") or "")[:16].replace("T", " ")
            col_a, col_b, col_c = st.columns([4, 2, 1])
            col_a.write(f"**{tpl.get('name', '—')}**")
            col_b.caption(f"{tpl.get('title', '—')} · {updated or '—'}")
            if col_c.button("🗑️", key=f"settings_del_tpl_{i}"):
                ok, msg = delete_template(tpl["id"])
                if ok:
                    st.success(msg)
                else:
                    st.error(msg)
                st.rerun()

    st.divider()
    st.subheader("📋 Список вакансий")
    vacancies = load_vacancies()
    if not vacancies:
        st.info("Нет вакансий. Создайте первую на вкладке «Вакансии».")
    else:
        for i, vac in enumerate(vacancies):
            status = "активна" if vac.get("active", True) else "в архиве"
            period = format_vacancy_search_period(vac)
            vac_id = vac.get("id") or f"idx_{i}"
            confirm_key = f"settings_confirm_del_vac_{vac_id}"
            is_confirming = bool(st.session_state.get(confirm_key))
            col_a, col_b, col_c, col_d = st.columns([4, 2, 1, 1])
            col_a.write(f"**{vac['title']}**")
            col_b.caption(f"{period} · {status}")
            if is_confirming:
                col_c.caption("⚠️ подтвердите")
                if col_d.button("Удалить", key=f"settings_del_vac_yes_{vac_id}", type="primary"):
                    from vacancy_store import delete_vacancy_by_id

                    ok, msg = delete_vacancy_by_id(vac.get("id"))
                    st.session_state.pop(confirm_key, None)
                    if ok:
                        st.success(msg)
                    else:
                        st.error(msg)
                    st.rerun()
                if col_b.button("Отмена", key=f"settings_del_vac_no_{vac_id}"):
                    st.session_state.pop(confirm_key, None)
                    st.rerun()
            else:
                col_c.caption(status)
                if col_d.button("🗑️", key=f"settings_del_vac_{vac_id}"):
                    st.session_state[confirm_key] = True
                    st.rerun()

