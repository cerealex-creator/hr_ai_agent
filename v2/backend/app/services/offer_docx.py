"""Generate candidate offer .docx from template + saved offer draft."""

from __future__ import annotations

import base64
import re
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.parse import quote

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from sqlalchemy.orm import Session

from app.core.config import Settings, get_settings
from app.db import models
from app.services.offer_draft import (
    company_logo_data_url,
    format_probation_months,
    get_offer_draft,
    prefill_offer_draft,
)

_TOKEN_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")


def default_template_path() -> Path:
    return Path(__file__).resolve().parent.parent / "assets" / "offer_template.docx"


def resolve_template_path(settings: Settings | None = None) -> Path:
    settings = settings or get_settings()
    raw = (getattr(settings, "offer_template_path", None) or "").strip()
    if raw:
        return Path(raw).expanduser()
    custom = custom_template_path(settings)
    if custom.is_file() and custom.stat().st_size > 0:
        return custom
    return default_template_path()


def custom_template_path(settings: Settings | None = None) -> Path:
    settings = settings or get_settings()
    return settings.resolved_legacy_data_dir() / "offer_template.docx"


def offer_template_info(settings: Settings | None = None) -> dict[str, Any]:
    settings = settings or get_settings()
    env_raw = (getattr(settings, "offer_template_path", None) or "").strip()
    if env_raw:
        p = Path(env_raw).expanduser()
        return {
            "source": "env",
            "path": str(p),
            "filename": p.name if p.is_file() else None,
            "has_custom": p.is_file(),
            "can_upload": False,
        }
    custom = custom_template_path(settings)
    default = default_template_path()
    if custom.is_file() and custom.stat().st_size > 0:
        return {
            "source": "upload",
            "path": str(custom),
            "filename": custom.name,
            "has_custom": True,
            "can_upload": True,
        }
    return {
        "source": "default",
        "path": str(default),
        "filename": default.name,
        "has_custom": False,
        "can_upload": True,
    }


def save_custom_offer_template(content: bytes, settings: Settings | None = None) -> Path:
    if not content or len(content) < 100:
        raise ValueError("Файл шаблона слишком маленький")
    if len(content) > 8_000_000:
        raise ValueError("Шаблон больше 8 МБ — упростите файл")
    if content[:2] != b"PK":
        raise ValueError("Нужен файл Word (.docx)")
    path = custom_template_path(settings)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(content)
    return path


def clear_custom_offer_template(settings: Settings | None = None) -> None:
    path = custom_template_path(settings)
    if path.is_file():
        path.unlink()


def ensure_default_template(path: Path | None = None) -> Path:
    """Create offer template matching HR letter structure with {{placeholders}}."""
    path = path or default_template_path()
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.is_file() and path.stat().st_size > 0:
        return path

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    doc.add_paragraph("{{greeting}} {{name_patronymic}}!")
    doc.add_paragraph(
        "От лица {{company}} сообщаем об успешном прохождении Вами собеседования "
        "на должность «{{position}}». Благодарим Вас за интерес, проявленный к нашей "
        "вакансии, и настоящим письмом подтверждаем намерение заключить с Вами "
        "трудовой договор на неопределённый срок."
    )
    doc.add_paragraph(
        "Рабочее место расположено по адресу {{office_address}}. "
        "Предлагаем {{work_schedule}}."
    )
    doc.add_paragraph(
        "Мы будем рады видеть Вас в нашем коллективе с {{start_date}}. "
        "Испытательный срок в целях проверки соответствия поручаемой работе — "
        "{{probation_months}}."
    )
    doc.add_paragraph("Наша компания готова предложить Вам следующие условия работы:")
    doc.add_paragraph(
        "Ежемесячная заработная плата на период испытательного срока – "
        "{{salary_probation_line}}.",
        style="List Paragraph",
    )
    doc.add_paragraph(
        "Заработная плата после испытательного срока – {{salary_after_line}}.",
        style="List Paragraph",
    )
    doc.add_paragraph(
        "Ежегодный отпуск в размере 28 календарных дней.",
        style="List Paragraph",
    )
    doc.add_paragraph(
        "40-часовая рабочая неделя с двумя выходными днями (суббота и воскресенье).",
        style="List Paragraph",
    )
    doc.add_paragraph(
        "Поддержка в период адаптации и гибкий подход к решению рабочих вопросов.",
        style="List Paragraph",
    )
    doc.add_paragraph(
        "Прочие условия работы будут оговорены в Вашем трудовом договоре.",
        style="List Paragraph",
    )
    doc.add_paragraph("")
    doc.add_paragraph("В ваши трудовые обязанности будут, в том числе, входить задачи:")
    doc.add_paragraph("{{duties}}")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Мы уверены, что данная должность в нашей компании даст Вам возможность "
        "ярко раскрыть свои опыт и талант, достичь высоких результатов!"
    )
    doc.add_paragraph(
        "В случае Вашего согласия с настоящим предложением о работе, просим Вас "
        "представить скан подписанного Вами оффера."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Руководитель. ______________________________ {{manager_name}}")
    doc.add_paragraph("")
    doc.add_paragraph(
        "С предложением согласна, намерение заключить трудовой договор подтверждаю "
        "________________________ {{full_name}}"
    )
    doc.save(path)
    return path


def build_offer_values(db: Session, candidate: models.Candidate) -> dict[str, str]:
    draft = get_offer_draft(candidate)
    if not any(draft.values()):
        draft = prefill_offer_draft(db, candidate)

    def dash(v: str) -> str:
        return v.strip() if (v or "").strip() else "—"

    values = {
        "greeting": dash(draft.get("greeting") or "Уважаемый(ая)"),
        "name_patronymic": dash(draft.get("name_patronymic") or draft.get("full_name") or ""),
        "full_name": dash(draft.get("full_name") or candidate.name or ""),
        "fio": dash(draft.get("full_name") or candidate.name or ""),
        "name": dash(draft.get("full_name") or candidate.name or ""),
        "company": dash(draft.get("company") or ""),
        "client": dash(draft.get("company") or ""),
        "position": dash(draft.get("position") or ""),
        "vacancy": dash(draft.get("position") or ""),
        "office_address": dash(draft.get("office_address") or ""),
        "work_schedule": dash(draft.get("work_schedule") or ""),
        "start_date": dash(draft.get("start_date") or ""),
        "probation_months": dash(format_probation_months(draft.get("probation_months") or "")),
        "salary_probation_line": dash(draft.get("salary_probation_line") or ""),
        "salary_after_line": dash(draft.get("salary_after_line") or ""),
        "salary": dash(draft.get("salary_after_line") or draft.get("salary_probation_line") or ""),
        "duties": (draft.get("duties") or "").strip() or "—",
        "manager_name": dash(draft.get("manager_name") or ""),
        "offer_date": dash(draft.get("start_date") or ""),
        "date": dash(draft.get("start_date") or ""),
    }
    return values


def _replace_text(text: str, values: dict[str, str]) -> str:
    def repl(m: re.Match[str]) -> str:
        key = m.group(1).strip().lower()
        if key in values:
            return values[key]
        return m.group(0)

    return _TOKEN_RE.sub(repl, text)


def _replace_in_paragraph(paragraph: Any, values: dict[str, str]) -> None:
    full = paragraph.text or ""
    if "{{" not in full:
        return
    new = _replace_text(full, values)
    if new == full:
        return
    if not paragraph.runs:
        paragraph.add_run(new)
        return
    paragraph.runs[0].text = new
    for run in paragraph.runs[1:]:
        run.text = ""


def _iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part is None:
                continue
            for p in part.paragraphs:
                yield p
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p


_BULLET_PREFIX_RE = re.compile(r"^[\u2022\u00b7•\-–—]\s*")


def _with_bullet_prefix(text: str) -> str:
    s = (text or "").strip()
    if not s:
        return s
    if _BULLET_PREFIX_RE.match(s):
        return s
    return f"• {s}"


def _set_paragraph_text(paragraph: Any, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _format_work_conditions_bullets(doc: Document) -> None:
    """Условия работы — тот же маркированный список, что у обязанностей."""
    in_section = False
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        low = text.lower()
        if "условия работы" in low:
            in_section = True
            continue
        if not in_section:
            continue
        if "трудовые обязанности" in low or "{{duties}}" in text:
            break
        if not text:
            continue
        _set_paragraph_text(paragraph, _with_bullet_prefix(text))


def _expand_duties_paragraphs(doc: Document, duties_text: str) -> None:
    lines = [_with_bullet_prefix(ln) for ln in (duties_text or "").splitlines() if ln.strip()]
    if not lines:
        lines = ["—"]
    for paragraph in list(doc.paragraphs):
        if "{{duties}}" not in (paragraph.text or "") and "{{ duties }}" not in (paragraph.text or ""):
            continue
        p_element = paragraph._element
        parent = p_element.getparent()
        idx = parent.index(p_element)
        _set_paragraph_text(paragraph, lines[0])
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph

        for j, line in enumerate(lines[1:], start=1):
            new_el = OxmlElement("w:p")
            parent.insert(idx + j, new_el)
            new_p = Paragraph(new_el, paragraph._parent)
            new_p.add_run(line)
            try:
                new_p.style = paragraph.style
            except Exception:
                pass
        return


def _insert_logo(doc: Document, data_url: str | None) -> None:
    if not data_url:
        return
    m = re.match(r"^data:image/(png|jpeg|jpg|gif|webp);base64,(.+)$", data_url, re.I | re.S)
    if not m:
        return
    raw = base64.b64decode(m.group(2))
    if len(raw) > 1_500_000:
        return
    section = doc.sections[0]
    header = section.header
    # Clear empty default para or add
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in list(paragraph.runs):
        run.text = ""
    run = paragraph.add_run()
    run.add_picture(BytesIO(raw), width=Cm(3.5))


def render_offer_docx(
    template_path: Path,
    values: dict[str, str],
    *,
    logo_data_url: str | None = None,
) -> bytes:
    if not template_path.is_file():
        raise FileNotFoundError(f"Шаблон оффера не найден: {template_path}")
    doc = Document(str(template_path))
    _expand_duties_paragraphs(doc, values.get("duties") or "")
    # After expansion, duties token is gone; still replace other tokens
    for p in _iter_all_paragraphs(doc):
        if "{{" in (p.text or ""):
            _replace_in_paragraph(p, values)
    _format_work_conditions_bullets(doc)
    _insert_logo(doc, logo_data_url)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def offer_filename(candidate: models.Candidate) -> str:
    raw = re.sub(r"[^\w\-]+", "_", (candidate.name or "candidate").strip(), flags=re.UNICODE)
    raw = raw.strip("_")[:60] or "candidate"
    return f"offer_{raw}.docx"


def attachment_content_disposition(filename: str) -> str:
    """ASCII fallback + UTF-8 filename for Starlette (headers must be latin-1)."""
    ascii_fallback = re.sub(r"[^A-Za-z0-9._-]+", "_", filename).strip("._-")
    if not ascii_fallback.lower().endswith(".docx"):
        ascii_fallback = f"{ascii_fallback or 'offer'}.docx"
    stem = ascii_fallback[:-5] if ascii_fallback.lower().endswith(".docx") else ascii_fallback
    if not re.search(r"[A-Za-z0-9]", stem):
        ascii_fallback = "offer.docx"
    encoded = quote(filename, safe="")
    return f"attachment; filename=\"{ascii_fallback}\"; filename*=UTF-8''{encoded}"


def generate_candidate_offer_docx(
    db: Session,
    candidate: models.Candidate,
    *,
    settings: Settings | None = None,
) -> tuple[bytes, str]:
    settings = settings or get_settings()
    path = resolve_template_path(settings)
    if path == default_template_path():
        ensure_default_template(path)
    elif not path.is_file():
        raise FileNotFoundError(
            f"Шаблон оффера не найден по пути OFFER_TEMPLATE_PATH={path}"
        )
    values = build_offer_values(db, candidate)
    vacancy = db.get(models.Vacancy, candidate.vacancy_id)
    logo = company_logo_data_url(db, vacancy)
    data = render_offer_docx(path, values, logo_data_url=logo)
    return data, offer_filename(candidate)
