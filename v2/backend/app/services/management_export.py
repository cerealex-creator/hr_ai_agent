"""СУП — ранний экспорт документов (цели / превью)."""
from __future__ import annotations

import html
import io
import re
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt
from sqlalchemy.orm import Session

from app.db import management_models as m
from app.services.management_system import goal_dimensions_map, list_goals, list_tasks


def _escape(s: str) -> str:
    return html.escape(s or "")


def build_goals_pack_html(
    db: Session,
    *,
    system: m.MgmtSystem,
    revision_id,
) -> str:
    goals = list_goals(db, revision_id)
    tasks = list_tasks(db, revision_id)
    dim_map = goal_dimensions_map(db, [g.id for g in goals])
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    kind_label = {"company": "Компания", "holding": "Холдинг", "demo": "Демо"}.get(system.kind, system.kind)

    rows = []
    for g in goals:
        dims = dim_map.get(g.id, [])
        dim_titles = ", ".join(d["title"] for d in dims) or "—"
        baseline = f"{g.baseline_value}" if g.baseline_value is not None else "—"
        target = f"{g.target_value}" if g.target_value is not None else "—"
        unit = g.metric_unit or ""
        rows.append(
            f"<tr><td>{_escape(g.title)}</td><td>{_escape(dim_titles)}</td>"
            f"<td>{_escape(baseline)} → {_escape(target)} {_escape(unit)}</td>"
            f"<td>{_escape(g.status)}</td></tr>"
        )

    task_rows = "".join(
        f"<li><strong>{_escape(t.title)}</strong> · {t.status}</li>" for t in tasks
    ) or "<li class='muted'>Задач пока нет</li>"

    coming = """
    <section class="soon">
      <h2>Скоро в экспорте</h2>
      <ul>
        <li>Оргсхема (PDF / PNG) — после шагов оргструктуры</li>
        <li>Должностные инструкции (DOCX / PDF) — после утверждения ролей (L3)</li>
        <li>Чек-листы и KPI по ролям</li>
      </ul>
    </section>
    """

    return f"""<!DOCTYPE html>
<html lang="ru">
<head>
<meta charset="utf-8"/>
<title>Пакет целей — {_escape(system.title)}</title>
<style>
  body {{ font-family: Georgia, 'Times New Roman', serif; max-width: 820px; margin: 2rem auto; padding: 0 1.25rem; color: #1a1a1a; }}
  h1 {{ font-size: 1.6rem; margin-bottom: 0.25rem; }}
  .meta {{ color: #666; font-size: 0.9rem; margin-bottom: 1.5rem; }}
  table {{ width: 100%; border-collapse: collapse; margin: 1rem 0; }}
  th, td {{ border: 1px solid #ccc; padding: 0.5rem 0.65rem; text-align: left; vertical-align: top; }}
  th {{ background: #f4f4f4; }}
  .soon {{ margin-top: 2rem; padding: 1rem; border: 1px dashed #aaa; background: #fafafa; }}
  .muted {{ color: #888; }}
  @media print {{
    body {{ margin: 0; }}
    .no-print {{ display: none; }}
    .soon {{ break-inside: avoid; }}
  }}
</style>
</head>
<body>
  <p class="no-print meta"><button onclick="window.print()">Печать / сохранить PDF</button></p>
  <h1>Пакет целей (L0)</h1>
  <p class="meta">{_escape(system.title)} · {kind_label} · черновик · {now}</p>
  <table>
    <thead><tr><th>Цель</th><th>Измерение</th><th>Сейчас → цель</th><th>Статус</th></tr></thead>
    <tbody>
      {''.join(rows) if rows else '<tr><td colspan="4" class="muted">Целей пока нет</td></tr>'}
    </tbody>
  </table>
  <h2>Задачи (L1)</h2>
  <ul>{task_rows}</ul>
  {coming}
</body>
</html>"""


def build_goals_pack_docx(
    db: Session,
    *,
    system: m.MgmtSystem,
    revision_id,
) -> bytes:
    goals = list_goals(db, revision_id)
    tasks = list_tasks(db, revision_id)
    dim_map = goal_dimensions_map(db, [g.id for g in goals])

    doc = Document()
    doc.add_heading(f"Пакет целей — {system.title}", level=1)
    kind_label = {"company": "Компания", "holding": "Холдинг", "demo": "Демо"}.get(system.kind, system.kind)
    p = doc.add_paragraph(f"{kind_label} · статус системы: {system.status}")
    p.runs[0].font.size = Pt(10)

    doc.add_heading("Цели (L0)", level=2)
    if not goals:
        doc.add_paragraph("Целей пока нет.")
    else:
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "Цель"
        hdr[1].text = "Измерение"
        hdr[2].text = "Сейчас → цель"
        hdr[3].text = "Статус"
        for g in goals:
            dims = dim_map.get(g.id, [])
            dim_titles = ", ".join(d["title"] for d in dims) or "—"
            baseline = f"{g.baseline_value}" if g.baseline_value is not None else "—"
            target = f"{g.target_value}" if g.target_value is not None else "—"
            unit = g.metric_unit or ""
            row = table.add_row().cells
            row[0].text = g.title
            row[1].text = dim_titles
            row[2].text = f"{baseline} → {target} {unit}".strip()
            row[3].text = g.status

    doc.add_heading("Задачи (L1)", level=2)
    if not tasks:
        doc.add_paragraph("Задач пока нет.")
    else:
        for t in tasks:
            doc.add_paragraph(f"{t.title} ({t.status})", style="List Bullet")

    doc.add_heading("Скоро в экспорте", level=2)
    for item in (
        "Оргсхема (PDF / PNG)",
        "Должностные инструкции (DOCX / PDF)",
        "Чек-листы и KPI по ролям",
    ):
        doc.add_paragraph(item, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def safe_filename(title: str, suffix: str) -> str:
    raw = re.sub(r"[^\w\-]+", "_", (title or "sistema").strip(), flags=re.UNICODE)
    raw = raw.strip("_")[:60] or "sistema"
    return f"{raw}_goals{suffix}"
