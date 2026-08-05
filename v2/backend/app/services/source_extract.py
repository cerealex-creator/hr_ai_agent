"""Extract text from HR source files (txt/pdf/docx/xlsx) and classify media."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

MEDIA_EXT = {".mp3", ".mp4", ".wav", ".webm", ".mkv", ".ogg", ".m4a", ".mov", ".aac", ".flac"}
DOC_EXT = {".txt", ".md", ".pdf", ".docx", ".xlsx", ".xls", ".csv"}


def classify_path(path: str | Path) -> str:
    ext = Path(path).suffix.lower()
    if ext in MEDIA_EXT:
        return "media"
    if ext in DOC_EXT:
        return "document"
    return "unknown"


def extract_text_from_bytes(filename: str, content: bytes) -> str:
    name = (filename or "file").lower()
    ext = Path(name).suffix.lower()
    if not content:
        return ""
    if ext in {".txt", ".md", ".csv"} or not ext:
        try:
            return content.decode("utf-8-sig")
        except UnicodeDecodeError:
            return content.decode("cp1251", errors="ignore")
    if ext == ".pdf":
        return _pdf(content)
    if ext == ".docx":
        return _docx(content)
    if ext in {".xlsx", ".xls"}:
        return _xlsx(content)
    # Fallback: try utf-8
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError:
        return ""


def extract_text_from_path(path: str | Path) -> str:
    p = Path(path)
    return extract_text_from_bytes(p.name, p.read_bytes())


def _pdf(content: bytes) -> str:
    if not content.lstrip().startswith(b"%PDF"):
        return ""
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("Нужен пакет pypdf") from exc
    reader = PdfReader(BytesIO(content))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            continue
    return "\n".join(parts).strip()


def _docx(content: bytes) -> str:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("Нужен пакет python-docx") from exc
    document = docx.Document(BytesIO(content))
    parts = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _xlsx(content: bytes) -> str:
    try:
        import openpyxl
    except ImportError as exc:
        raise RuntimeError("Нужен пакет openpyxl") from exc
    wb = openpyxl.load_workbook(BytesIO(content), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"## {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()
